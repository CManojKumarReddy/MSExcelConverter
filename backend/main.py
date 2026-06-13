"""
DocToExcel - FastAPI Backend
Converts PDF, DOCX, Images, CSV, TXT to Excel (.xlsx)
"""

import os
import re as _re
import time
import uuid
import asyncio
import logging
import traceback
from pathlib import Path
from typing import Optional
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger(__name__)

# ── Optional config (.env) ─────────────────────────────────────────────────────
# Load a local .env if python-dotenv is installed (it's optional — the server
# runs fine without it; cloud OCR simply stays disabled).
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env")
except Exception:
    pass

AZURE_DI_ENDPOINT = os.getenv("AZURE_DI_ENDPOINT", "").strip()
AZURE_DI_KEY      = os.getenv("AZURE_DI_KEY", "").strip()
GEMINI_API_KEY    = os.getenv("GEMINI_API_KEY", "").strip()


def _azure_configured() -> bool:
    """True only when both the Azure Document Intelligence endpoint and key are set."""
    return bool(AZURE_DI_ENDPOINT and AZURE_DI_KEY)


def _gemini_configured() -> bool:
    """True when a Google Gemini (AI Studio) API key is set."""
    return bool(GEMINI_API_KEY)


def _cloud_ocr_engine() -> Optional[str]:
    """Which cloud OCR engine is available: 'gemini' (preferred) > 'azure' > None."""
    if _gemini_configured():
        return "gemini"
    if _azure_configured():
        return "azure"
    return None


class CloudOCRUnavailable(Exception):
    """Raised when a cloud OCR path can't run (not installed / not configured / API error)."""

# ── App Setup ─────────────────────────────────────────────────────────────────
app = FastAPI(title="DocToExcel API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUTS_DIR = Path(__file__).parent / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".csv", ".txt"}

# ── Conversion Job Queue ──────────────────────────────────────────────────────
#
# Conversions are CPU/RAM-heavy (Tesseract, pdfplumber, openpyxl). Running them
# all at once under load exhausts memory and crashes the box. So every request
# is turned into a Job and admitted through a semaphore that caps how many run
# at the same time; the rest wait in a FIFO queue.
#
#   - MAX_CONCURRENT_CONVERSIONS   how many conversions run simultaneously.
#                                  Bump this when you upsize the host.
#   - JOB_RETENTION_SECONDS        how long finished jobs linger before pruning
#                                  (so the client has time to read the result).
#
# A queued job can be cancelled (removed from the queue). A job that has already
# started processing cannot — a running conversion in a worker thread can't be
# safely interrupted.

MAX_CONCURRENT = max(1, int(os.getenv("MAX_CONCURRENT_CONVERSIONS", "2")))
JOB_RETENTION_SECONDS = int(os.getenv("JOB_RETENTION_SECONDS", "600"))

_executor = ThreadPoolExecutor(max_workers=MAX_CONCURRENT, thread_name_prefix="convert")
_semaphore: Optional[asyncio.Semaphore] = None
_job_seq = 0


def _get_semaphore() -> asyncio.Semaphore:
    """Create the concurrency semaphore lazily, bound to the running event loop."""
    global _semaphore
    if _semaphore is None:
        _semaphore = asyncio.Semaphore(MAX_CONCURRENT)
    return _semaphore


class Job:
    """A single queued/running conversion and its result."""

    def __init__(self, params: dict):
        global _job_seq
        _job_seq += 1
        self.seq = _job_seq
        self.id = uuid.uuid4().hex
        self.params = params                 # args for _run_conversion
        self.status = "queued"               # queued|processing|done|error|cancelled|password_required
        self.output_filename: Optional[str] = None
        self.message: Optional[str] = None
        self.error: Optional[str] = None
        self.password_provided = bool(params.get("password"))
        self.created_at = time.time()
        self.finished_at: Optional[float] = None
        self.task: Optional[asyncio.Task] = None


JOBS: dict[str, Job] = {}


def _prune_jobs():
    """Drop finished jobs older than the retention window to bound memory."""
    now = time.time()
    stale = [
        jid for jid, j in JOBS.items()
        if j.finished_at is not None and now - j.finished_at > JOB_RETENTION_SECONDS
    ]
    for jid in stale:
        JOBS.pop(jid, None)


def _queue_position(job: Job) -> int:
    """How many not-yet-finished jobs sit ahead of this one (0 = up next/running)."""
    return sum(
        1 for j in JOBS.values()
        if j.seq < job.seq and j.status in ("queued", "processing")
    )


async def _process_job(job: Job):
    """Wait for a free slot, then run the conversion in a worker thread."""
    sem = _get_semaphore()
    try:
        async with sem:
            # The job may have been cancelled while waiting for a slot.
            if job.status == "cancelled":
                return
            job.status = "processing"
            loop = asyncio.get_event_loop()
            try:
                output_filename, msg = await loop.run_in_executor(
                    _executor, _run_conversion, job.params
                )
                job.output_filename = output_filename
                job.message = msg
                job.status = "done"
            except HTTPException as e:
                if e.status_code == 423:          # password-protected PDF
                    job.status = "password_required"
                else:
                    job.status = "error"
                    job.error = str(e.detail)
            except Exception as e:
                log.error("Conversion error:\n%s", traceback.format_exc())
                job.status = "error"
                job.error = f"Conversion failed: {e}"
    except asyncio.CancelledError:
        # Cancelled while still waiting in the queue.
        job.status = "cancelled"
        raise
    finally:
        if job.status in ("done", "error", "cancelled", "password_required"):
            job.finished_at = time.time()


def _job_view(job: Job) -> dict:
    """Serialise a job for the client."""
    view = {"job_id": job.id, "status": job.status}
    if job.status == "queued":
        view["position"] = _queue_position(job)
    elif job.status == "done":
        view["output_filename"] = job.output_filename
        view["message"] = job.message
    elif job.status == "error":
        view["detail"] = job.error or "Conversion failed."
    elif job.status == "password_required":
        view["password_provided"] = job.password_provided
    return view


# ── Excel Helpers ─────────────────────────────────────────────────────────────

def style_header_row(ws, row_num: int, num_cols: int):
    """Apply teal header styling to a row."""
    header_fill = PatternFill(start_color="009E76", end_color="009E76", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row_num, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def apply_table_borders(ws, min_row: int, max_row: int, min_col: int, max_col: int):
    """Apply thin borders to a cell range."""
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
        for cell in row:
            cell.border = border


_INVALID_SHEET_CHARS = str.maketrans({c: "" for c in r"\/*?:[]\'"})

def sanitize_sheet_name(name: str) -> str:
    """Remove Excel-invalid characters and truncate to 31 chars."""
    return name.translate(_INVALID_SHEET_CHARS).strip()[:31] or "Sheet"




# ── Stage 4: Merge continuation rows ──────────────────────────────────────────

def _merge_continuation_rows(rows: list[list[str]], col_boundaries: list[int] | None = None) -> list[list[str]]:
    """
    Bank statement descriptions often wrap across multiple visual lines, producing
    extra rows with content only in the Description column.

    A row is a *continuation* of the preceding row when:
    - No anchor column (columns to the left of the widest/description column)
      contains a pure date  →  it is not a new transaction
    - No financial-column cell (columns to the right of description) contains an
      amount  →  it carries no new financial data

    The description column is inferred as the widest column by pixel width.
    """
    if not rows:
        return rows

    if col_boundaries and len(col_boundaries) > 2:
        widths   = [col_boundaries[i + 1] - col_boundaries[i]
                    for i in range(len(col_boundaries) - 1)]
        desc_col = widths.index(max(widths))
    else:
        desc_col = min(1, len(rows[0]) - 1)

    anchor_cols = list(range(desc_col))
    ncols_hint  = max(len(rows[0]), desc_col + 2)

    merged: list[list[str]] = [list(rows[0])]

    for row in rows[1:]:
        # Skip empty rows
        if not any(c.strip() for c in row):
            continue

        # A complete header row is always a new row — never a continuation of
        # whatever came before (e.g. a MICR/address metadata row).
        if _is_header_row(row):
            merged.append(list(row))
            continue

        # Skip stray header-fragment rows — lines from multi-line column headers
        # (e.g. lone "Date" sitting below "Value" in a "Value Date" cell) that
        # contain only header keywords but do NOT form a complete header row.
        row_words = [w for c in row for w in c.strip().split() if c.strip()]
        if row_words and all(w.lower() in _HEADER_KEYWORDS for w in row_words):
            continue

        anchor_has_date = any(
            bool(_PURE_DATE_RE.match(row[i].strip()))
            for i in anchor_cols
            if i < len(row) and row[i].strip()
        )
        has_financial = any(
            _NUMERIC_RE.match(row[i].strip())
            for i in range(desc_col + 1, max(len(row), ncols_hint))
            if i < len(row) and row[i].strip()
        )

        if not anchor_has_date and not has_financial:
            prev = merged[-1]
            while len(prev) <= desc_col:
                prev.append("")
            text = " ".join(c.strip() for c in row if c.strip())
            if text:
                prev[desc_col] = (prev[desc_col] + " " + text).strip() if prev[desc_col] else text
        else:
            merged.append(list(row))

    return merged


# ── Stage 5: Build single-sheet output ────────────────────────────────────────

def _is_header_row(row: list[str]) -> bool:
    """
    Return True if this row is a table column-header line.

    Strict criteria:
    - Every non-empty cell must be ≤ 3 words (real headers are concise labels)
    - Word set must hit both required keyword groups
    """
    non_empty = [c.strip() for c in row if c.strip()]
    if not non_empty:
        return False
    if any(len(c.split()) > 3 for c in non_empty):
        return False        # long cell → address / metadata, not a header
    texts = {w.lower() for c in non_empty for w in c.split()}
    return all(texts & grp for grp in _HEADER_REQUIRED_GROUPS)


def _build_single_sheet_rows(all_sections: list[tuple[str, list[list[str]]]]) -> list[list[str]]:
    """
    Combine all sections into one sheet:

    - For sections that contain a header row: skip everything before the header
      (address text, MICR lines, metadata), emit the header exactly once globally,
      then emit all subsequent data rows.
    - For sections without a header row: include only if they contain financial
      amounts (skips pure-text metadata blocks that contain dates like "16 Jul 2019").
    - Duplicate header rows from subsequent pages are silently dropped.
    """
    combined: list[list[str]] = []
    header_emitted = False

    for _name, rows in all_sections:
        if not rows:
            continue

        header_idx = next((i for i, row in enumerate(rows) if _is_header_row(row)), None)

        if header_idx is not None:
            if not header_emitted:
                header_row = list(rows[header_idx])
                # Absorb any immediately-following header-keyword-only rows into
                # the header.  This handles multi-line column headers in the PDF
                # (e.g. "Value" on line 1, "Date" on line 2 of the same cell).
                start_data = header_idx + 1
                while start_data < len(rows):
                    frag = rows[start_data]
                    frag_words = [w for c in frag for w in c.strip().split() if c.strip()]
                    if frag_words and all(w.lower() in _HEADER_KEYWORDS for w in frag_words):
                        # Merge fragment words into the header cell that covers
                        # the same column position.
                        for ci, cell in enumerate(frag):
                            if cell.strip() and ci < len(header_row):
                                header_row[ci] = (
                                    (header_row[ci] + " " + cell.strip()).strip()
                                    if header_row[ci] else cell.strip()
                                )
                        start_data += 1
                    else:
                        break
                combined.append(header_row)
                header_emitted = True
            else:
                start_data = header_idx + 1
            for row in rows[start_data:]:
                if not _is_header_row(row):
                    # Skip stray header-fragment rows (incomplete keyword-only lines)
                    row_words = [w for c in row for w in c.strip().split() if c.strip()]
                    if (row_words
                            and all(w.lower() in _HEADER_KEYWORDS for w in row_words)
                            and not _is_header_row(row)):
                        continue
                    combined.append(row)
        else:
            has_financial = any(
                _NUMERIC_RE.match(cell.strip())
                for row in rows for cell in row if cell.strip()
            )
            if not has_financial:
                continue
            for row in rows:
                # Skip pure metadata rows (no date, no financial amount)
                row_has_amount = any(
                    _NUMERIC_RE.match(c.strip()) for c in row if c.strip()
                )
                row_has_date = any(
                    _PURE_DATE_RE.match(c.strip()) for c in row if c.strip()
                )
                if row_has_amount or row_has_date:
                    combined.append(row)

    return combined


def auto_fit_columns(ws):
    """Set reasonable column widths based on content."""
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        adjusted = min(max(max_len + 4, 10), 60)
        ws.column_dimensions[col_letter].width = adjusted


def write_data_to_sheet(ws, rows: list[list], sheet_title: str = "Converted Data"):
    """Write a list-of-lists to a worksheet with nice formatting."""
    ws.title = sheet_title[:31]  # Excel sheet name limit

    if not rows:
        ws["A1"] = "No data found in the document."
        return

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.alignment = Alignment(vertical="center", wrap_text=False)

    # Style first row as header if it looks like one
    if len(rows) > 1:
        style_header_row(ws, 1, len(rows[0]))

    apply_table_borders(ws, 1, len(rows), 1, max(len(r) for r in rows))
    auto_fit_columns(ws)

    # Freeze the header row
    ws.freeze_panes = "A2"


def save_workbook(wb: openpyxl.Workbook, stem: str) -> str:
    """Save workbook to outputs dir and return just the filename."""
    unique_id = uuid.uuid4().hex[:8]
    filename = f"{stem}_{unique_id}.xlsx"
    out_path = OUTPUTS_DIR / filename
    wb.save(out_path)
    log.info("Saved: %s", out_path)
    return filename


# ── Converters ────────────────────────────────────────────────────────────────

def convert_csv(content: bytes, stem: str) -> str:
    """CSV → Excel. Detects delimiter automatically."""
    import csv
    import io

    text = content.decode("utf-8-sig", errors="replace")
    # Sniff delimiter
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(io.StringIO(text), dialect)
    rows = [row for row in reader if any(cell.strip() for cell in row)]

    wb = openpyxl.Workbook()
    ws = wb.active
    write_data_to_sheet(ws, rows, "CSV Data")
    return save_workbook(wb, stem)


def convert_txt(content: bytes, stem: str) -> str:
    """TXT → Excel. Tries tab-separated first, then splits into lines."""
    import csv
    import io

    text = content.decode("utf-8-sig", errors="replace")
    lines = [l for l in text.splitlines() if l.strip()]

    # Detect if tab-separated
    if lines and "\t" in lines[0]:
        reader = csv.reader(io.StringIO(text), delimiter="\t")
        rows = [row for row in reader if any(cell.strip() for cell in row)]
    else:
        # Put each line as a row with a single column
        rows = [["Line", "Content"]]
        for i, line in enumerate(lines, start=1):
            rows.append([i, line])

    wb = openpyxl.Workbook()
    ws = wb.active
    write_data_to_sheet(ws, rows, "Text Data")
    return save_workbook(wb, stem)


def _setup_tesseract():
    """Ensure pytesseract can find the Tesseract binary."""
    import shutil
    import pytesseract
    if not shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ── PDF Table Extraction Pipeline ────────────────────────────────────────────
#
# Architecture overview:
#   1. GROUP words into visual lines using a y-band approach (not a fixed tolerance)
#   2. SPLIT lines into sections at large vertical gaps
#   3. DETECT header row by keyword density — only accept lines where the majority
#      of words are header keywords (rejects address/MICR lines that happen to end
#      with "Deposit Withdrawal Balance")
#   4. CLUSTER multi-word column headers (e.g. "Value Date") by gap size, then place
#      column boundaries at cluster midpoints — not between every word
#   5. ASSIGN words to columns; pull stray year tokens back to adjacent date column
#   6. MERGE continuation rows (multiline descriptions) into their anchor row
#   7. FILTER metadata sections from single-sheet output

_HEADER_KEYWORDS = {
    "date", "value", "description", "debit", "credit", "balance", "amount",
    "withdrawal", "deposit", "cheque", "chq", "narration", "particulars",
    "transaction", "reference", "remarks",
}
_HEADER_REQUIRED_GROUPS = [
    {"debit", "credit", "amount", "withdrawal", "deposit"},
    {"date", "description", "balance", "narration", "particulars"},
]
# Known multi-word bank column headers — consecutive words matching one of
# these ordered bigrams are kept in the same cluster (single column).
_HEADER_BIGRAMS = {
    ("value", "date"),
    ("debit", "amount"),
    ("credit", "amount"),
    ("cheque", "no"),
    ("cheque", "no."),
    ("cheque", "number"),
    ("chq", "no"),
    ("chq", "no."),
    ("chq", "number"),
    ("reference", "no"),
    ("reference", "no."),
    ("reference", "number"),
    ("transaction", "date"),
    ("posting", "date"),
    ("sl", "no"),
    ("sl", "no."),
    ("s", "no"),
    ("s.", "no"),
}

_NUMERIC_RE      = _re.compile(r"^[\$\-]?\d{1,10}(,\d{3})*\.\d{2}$")
_PURE_DATE_RE    = _re.compile(
    r'^(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}'
    r'|\d{1,2}\s+[A-Za-z]{3,9}\.?\s+\d{2,4}'
    r'|[A-Za-z]{3,9}\.?\s+\d{1,2},?\s+\d{2,4}'
    r')$'
)
_PARTIAL_DATE_RE = _re.compile(r'\d{1,2}\s+[A-Za-z]{3,9}\.?\s*$')
_LONE_YEAR_RE    = _re.compile(r'^\d{2}$|^\d{4}$')


# ── Stage 1: Group words into visual lines ─────────────────────────────────────

def _group_into_lines(words: list[dict]) -> list[list[dict]]:
    """
    Group words into visual lines using bounding-box overlap, NOT y-tolerance.

    Algorithm:
    - Sort words by top coordinate.
    - Maintain the bottom edge of the current line (max bottom seen so far).
    - A new word joins the current line when its top edge is BELOW the current
      line's bottom edge minus a small allowance (25 % of avg height). This
      means words must genuinely OVERLAP the current line's vertical extent.
    - Otherwise a new line starts.

    Why this is better than tolerance-based approaches:
    - Two consecutive table rows whose bounding boxes DON'T overlap are always
      placed in separate line-groups, regardless of how close they are.
    - Words on the same visual line that have slightly different top values
      (common in pdfplumber output) are still grouped together because their
      boxes overlap.
    - The MICR line and the table header line are NEVER merged even when they
      are only 1-2 pt apart, because the MICR box ends before the header box
      begins.
    """
    if not words:
        return []

    words = sorted(words, key=lambda w: (w["top"], w["left"]))

    lines:      list[list[dict]] = []
    current     = [words[0]]
    line_bottom = words[0]["top"] + max(words[0]["height"], 1)

    for word in words[1:]:
        avg_h     = sum(w["height"] for w in current) / len(current)
        # Allow a small downward overlap tolerance (25 % of avg height) to
        # handle words on the same line whose tops differ slightly.
        overlap_threshold = line_bottom - avg_h * 0.25

        if word["top"] < overlap_threshold:
            # Word's top is above the current line's adjusted bottom → same line
            current.append(word)
            line_bottom = max(line_bottom, word["top"] + max(word["height"], 1))
        else:
            lines.append(sorted(current, key=lambda w: w["left"]))
            current     = [word]
            line_bottom = word["top"] + max(word["height"], 1)

    lines.append(sorted(current, key=lambda w: w["left"]))
    return lines


# ── Stage 2: Detect column boundaries from the header line ────────────────────

def _find_col_boundaries(line_groups: list[list[dict]], img_width: int) -> list[int] | None:
    """
    Scan line_groups for the table header row and return column-boundary x-positions.

    A genuine header line:
    - contains words from BOTH required keyword groups
    - has at least 3 words
    - has ≥ 50 % of its words as header keywords (rejects MICR/address lines that
      accidentally end with "… Deposit Withdrawal Balance Date")

    Multi-word column headers ("Value Date", "Chq No") are detected by matching
    consecutive word pairs against a known-bigrams list.  Boundaries are placed at
    the midpoint between adjacent clusters, not between individual words.
    """
    # Font-based header detection (lesson from professional converters):
    # pdfplumber words carry 'fontname' and 'size'. Bold or larger-than-median
    # text is almost certainly a header row even with fewer keyword matches.
    all_sizes   = [w.get("size", 0) or 0 for line in line_groups for w in line]
    median_size = sorted(all_sizes)[len(all_sizes) // 2] if all_sizes else 0.0

    def _is_bold_line(words: list[dict]) -> bool:
        sizes = [w.get("size", 0) or 0 for w in words]
        names = [w.get("fontname", "") or "" for w in words]
        larger = sum(1 for s in sizes if s > median_size * 1.1) > len(sizes) * 0.5
        bold   = any(tok in n for n in names for tok in ("Bold", "Heavy", "Black", "Demi"))
        return larger or bold

    for line in line_groups:
        texts    = {w["text"].lower() for w in line}
        is_bold  = _is_bold_line(line)

        # Bold lines need only 1 keyword hit; non-bold need both required groups.
        if is_bold:
            if len(texts & _HEADER_KEYWORDS) < 1 or len(line) < 2:
                continue
        else:
            if not (all(texts & grp for grp in _HEADER_REQUIRED_GROUPS) and len(line) >= 3):
                continue
            kw_count = sum(1 for w in line if w["text"].lower() in _HEADER_KEYWORDS)
            if kw_count < len(line) / 2:
                continue      # mostly non-keyword words → address/metadata line

        header_words = sorted(line, key=lambda w: w["left"])

        # Cluster consecutive words that form known multi-word column headers
        # (e.g. "Value" + "Date" → "Value Date").  Any other consecutive pair
        # starts a new cluster — gap size is NOT used (too unreliable across PDFs).
        clusters: list[list[dict]] = [[header_words[0]]]
        for word in header_words[1:]:
            bigram = (clusters[-1][-1]["text"].lower(), word["text"].lower())
            if bigram in _HEADER_BIGRAMS:
                clusters[-1].append(word)
            else:
                clusters.append([word])

        # Place boundaries at the midpoint between adjacent clusters
        boundaries = [0]          # start from page left edge (not header word edge)
        for i in range(len(clusters) - 1):
            right_cur  = max(w["right"] for w in clusters[i])
            left_next  = clusters[i + 1][0]["left"]
            boundaries.append((right_cur + left_next) // 2)
        boundaries.append(img_width)
        return boundaries

    return None


# ── Stage 3: Assign words to columns ──────────────────────────────────────────

def _line_to_cells_fixed(line: list[dict], col_boundaries: list[int]) -> list[str]:
    """
    Map each word to a column bucket using col_boundaries.

    Special rules:
    - The rightmost column (Balance) only accepts numbers.  Non-numeric words that
      drift right are folded back into the Description column (second-to-last).
    - A lone 2- or 4-digit year (e.g. "19", "2019") that drifts one column to the
      right of a partial date ("16 Jun") is pulled back into the date's column.
    """
    num_cols = len(col_boundaries) - 1
    if num_cols < 1:
        return [" ".join(w["text"] for w in line)]

    cells    = [""] * num_cols
    last     = num_cols - 1
    prev_col: int | None = None

    for word in line:
        # Find which column bucket this word's left edge falls into
        col_idx = 0
        for j in range(1, num_cols):
            if word["left"] >= col_boundaries[j]:
                col_idx = j

        # Rightmost column guard: only numerics and header keywords allowed
        if (col_idx == last and num_cols > 1
                and not _NUMERIC_RE.match(word["text"])
                and word["text"].lower() not in _HEADER_KEYWORDS):
            col_idx = last - 1

        # Lone-year pull-back: "16 Jun" | "19" → "16 Jun 19" | ""
        if (prev_col is not None
                and col_idx != prev_col
                and _LONE_YEAR_RE.match(word["text"])
                and _PARTIAL_DATE_RE.search(cells[prev_col])):
            col_idx = prev_col

        cells[col_idx] = (cells[col_idx] + " " + word["text"]).strip()
        prev_col = col_idx

    return cells


def _line_to_cells_gap(line: list[dict]) -> list[str]:
    """Split a line into cells by detecting large horizontal gaps (fallback)."""
    cells: list[str] = []
    bucket = [line[0]]
    for word in line[1:]:
        prev         = bucket[-1]
        gap          = word["left"] - prev["right"]
        avg_char_w   = prev["width"] / max(len(prev["text"]), 1)
        if gap > avg_char_w * 2.5:
            cells.append(" ".join(w["text"] for w in bucket))
            bucket = [word]
        else:
            bucket.append(word)
    cells.append(" ".join(w["text"] for w in bucket))
    return cells


def _ocr_image_to_sections(img) -> list[tuple[str, list[list[str]]]]:
    """
    OCR an image and split into sections by large vertical gaps between line groups.
    Each section gets its own column boundaries (derived from its header row).
    Returns [(section_title, rows), ...].
    """
    import pytesseract

    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    words = []
    for i in range(len(data["text"])):
        text = data["text"][i].strip()
        if text and int(data["conf"][i]) > 0:
            words.append({
                "text": text,
                "left": data["left"][i],
                "top": data["top"][i],
                "width": data["width"][i],
                "height": data["height"][i],
                "right": data["left"][i] + data["width"][i],
            })

    if not words:
        return []

    line_groups = _group_into_lines(words)
    img_width = getattr(img, "width", 9999)

    # Determine threshold for a "section break" gap (2.5× the median line spacing)
    if len(line_groups) > 1:
        tops = [line[0]["top"] for line in line_groups]
        spacings = sorted([tops[i + 1] - tops[i] for i in range(len(tops) - 1)])
        median_spacing = spacings[len(spacings) // 2]
        gap_threshold = median_spacing * 2.5
    else:
        gap_threshold = 9999

    # Split line_groups into section buckets at large vertical gaps
    section_buckets: list[list[list[dict]]] = [[]]
    for i, line in enumerate(line_groups):
        section_buckets[-1].append(line)
        if i + 1 < len(line_groups):
            gap = line_groups[i + 1][0]["top"] - line[0]["top"]
            if gap > gap_threshold:
                section_buckets.append([])

    # Convert each bucket to rows with per-section column detection
    sections: list[tuple[str, list[list[str]]]] = []
    for bucket in section_buckets:
        if not bucket:
            continue
        col_boundaries = _find_col_boundaries(bucket, img_width)
        rows: list[list[str]] = []
        for line in bucket:
            cells = _line_to_cells_fixed(line, col_boundaries) if col_boundaries else _line_to_cells_gap(line)
            if any(c.strip() for c in cells):
                rows.append(cells)
        if rows:
            rows = _merge_continuation_rows(rows, col_boundaries)
            title = sanitize_sheet_name(" ".join(c for c in rows[0] if c.strip())) or "Section"
            sections.append((title, rows))

    return sections


def _ocr_pdf_pages(content: bytes) -> list[tuple[str, list[list[str]]]]:
    """
    Render each PDF page to an image, OCR it, and split by vertical gaps.
    Returns [(sheet_name, rows), ...] — one entry per section across all pages.
    """
    try:
        import pytesseract
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="pytesseract not installed. Run: pip install pytesseract  (and install Tesseract OCR)"
        )
    _setup_tesseract()

    try:
        import pdfplumber
    except ImportError:
        raise HTTPException(status_code=500, detail="pdfplumber is not installed.")

    import io

    results: list[tuple[str, list[list[str]]]] = []
    used_names: dict[str, int] = {}

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            img = page.to_image(resolution=200).original
            for title, rows in _ocr_image_to_sections(img):
                base = f"P{page_num} {title}"[:31]
                n = used_names.get(base, 0)
                used_names[base] = n + 1
                sheet_name = base if n == 0 else f"{base[:28]} {n}"
                results.append((sheet_name[:31], rows))

    return results


def convert_pdf(content: bytes, stem: str, mode: str = "single", password: str = "") -> str:
    """PDF → Excel. mode='separate' puts each table/section on its own sheet;
    mode='single' stacks everything into one sheet.
    Raises HTTP 423 if the PDF is password-protected and no/wrong password given."""
    try:
        import pdfplumber
    except ImportError:
        raise HTTPException(status_code=500, detail="pdfplumber is not installed. Run: pip install pdfplumber")

    import io
    from pdfminer.pdfdocument import PDFPasswordIncorrect
    from pdfminer.pdfparser import PDFSyntaxError

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # Collect all (sheet_name, rows) sections first
    all_sections: list[tuple[str, list[list[str]]]] = []

    open_kwargs = {"password": password} if password else {}
    try:
        pdf_file = pdfplumber.open(io.BytesIO(content), **open_kwargs)
    except (PDFPasswordIncorrect, PDFSyntaxError, Exception) as e:
        err = str(e).lower()
        if "password" in err or "encrypt" in err or "incorrect" in err or isinstance(e, PDFPasswordIncorrect):
            raise HTTPException(status_code=423, detail="password_required")
        raise

    # ── Two-pass word-based extraction ────────────────────────────────────────
    #
    # Pass 1 — collect every page's words and line_groups; find the BEST column
    #          boundaries across the entire document (the cleanest header row may
    #          be on page 2 even though we need it for page 1).
    #
    # Pass 2 — convert each page's line_groups into rows, using per-section
    #          boundaries when available, falling back to the document-level
    #          best boundaries otherwise.  This solves the MICR-merges-with-header
    #          problem: page 1's header may be obscured but page 2's is clean.

    def _normalise_words(raw_words):
        return [
            {
                "text":   w["text"].strip(),
                "left":   int(w["x0"]),
                "top":    int(w["top"]),
                "width":  int(w["x1"] - w["x0"]),
                "height": int(w["bottom"] - w["top"]),
                "right":  int(w["x1"]),
            }
            for w in raw_words if w["text"].strip()
        ]

    def _section_buckets(line_groups):
        """Split line_groups into section buckets at large vertical gaps."""
        if len(line_groups) > 1:
            tops = [ln[0]["top"] for ln in line_groups]
            spacings = sorted([tops[i + 1] - tops[i] for i in range(len(tops) - 1)])
            gap_threshold = spacings[len(spacings) // 2] * 2.5
        else:
            gap_threshold = 9999
        buckets: list[list[list[dict]]] = [[]]
        for i, line in enumerate(line_groups):
            buckets[-1].append(line)
            if i + 1 < len(line_groups):
                if line_groups[i + 1][0]["top"] - line[0]["top"] > gap_threshold:
                    buckets.append([])
        return buckets

    def _lines_to_rows(line_groups, col_boundaries):
        """Convert a list of line_groups to rows using col_boundaries."""
        rows: list[list[str]] = []
        for line in line_groups:
            if len(line) <= 2:
                merged = " ".join(w["text"] for w in line).strip()
                if merged:
                    rows.append([merged])
            else:
                cells = (
                    _line_to_cells_fixed(line, col_boundaries)
                    if col_boundaries else _line_to_cells_gap(line)
                )
                if any(c.strip() for c in cells):
                    rows.append(cells)
        return rows

    # ── Single pass: read words + detect boundaries + build sections ─────────
    # We need one pass because pdfplumber closes the PDF on context exit.
    # Strategy:
    #   1. Extract all page word data.
    #   2. Find the best (most-column) boundaries from any page.
    #   3. Process every section bucket, using per-section boundaries when
    #      available, falling back to the document-best when not (e.g. page 1
    #      where MICR text may prevent local header detection).

    best_boundaries: list[int] | None = None
    page_data: list[tuple[int, int, list]] = []  # (page_num, img_width, line_groups)
    # structured_sections: tables extracted via pdfplumber's grid detector (bypasses _merge_continuation_rows)
    structured_sections: list[tuple[str, list[list[str]]]] = []

    def _is_useful_table(t: list[list]) -> bool:
        """Return True if pdfplumber table has ≥2 columns and ≥2 data rows."""
        if not t or len(t) < 2:
            return False
        col_count = max((len(r) for r in t), default=0)
        if col_count < 2:
            return False
        non_empty_rows = sum(1 for r in t if any(c and str(c).strip() for c in r))
        return non_empty_rows >= 2

    def _is_degenerate_table(t: list[list]) -> bool:
        """
        Detect a pdfplumber table whose ROW segmentation failed: a bank statement
        with column lines but NO horizontal lines between transaction rows gets
        an entire page crammed into one row, with each cell holding many
        newline-separated values.  Signal: some cell holds far more text lines
        than the table has rows.  Such pages must fall back to the word-based
        banking pipeline (which segments rows correctly).
        """
        rows = len(t)
        max_cell_lines = 0
        for r in t:
            for c in r:
                if c:
                    max_cell_lines = max(max_cell_lines, str(c).count("\n") + 1)
        return max_cell_lines >= 4 and max_cell_lines > rows

    with pdf_file as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # ── Try structured table extraction first (bordered PDFs) ──────────
            try:
                tables = page.extract_tables()
            except Exception:
                tables = []
            page_tables = [t for t in (tables or []) if _is_useful_table(t)]
            # If any extracted table is degenerate (rows not segmented — typical
            # of bank statements with no inter-row lines), abandon the structured
            # path for the WHOLE page and let the word-based pipeline handle it.
            if page_tables and any(_is_degenerate_table(t) for t in page_tables):
                page_tables = []
            if page_tables:
                for tbl_idx, tbl in enumerate(page_tables, start=1):
                    # Normalise cells: None → ""
                    rows = [[str(c).strip() if c is not None else "" for c in r] for r in tbl]
                    rows = [r for r in rows if any(c for c in r)]
                    if rows:
                        title = sanitize_sheet_name(
                            " ".join(c for c in rows[0] if c) or f"Page {page_num} T{tbl_idx}"
                        )
                        structured_sections.append((title, rows))
                # Skip word-based extraction for this page — grid handled it
                page_data.append((page_num, int(page.width), []))
                continue

            raw = page.extract_words(keep_blank_chars=False, use_text_flow=False)
            if not raw:
                page_data.append((page_num, int(page.width), []))
                continue
            words       = _normalise_words(raw)
            img_width   = int(page.width)
            line_groups = _group_into_lines(words)
            page_data.append((page_num, img_width, line_groups))

            # Track the cleanest header found so far (most columns wins)
            boundaries = _find_col_boundaries(line_groups, img_width)
            if boundaries and (
                best_boundaries is None or len(boundaries) > len(best_boundaries)
            ):
                best_boundaries = boundaries

    # Structured tables go first (no _merge_continuation_rows needed)
    all_sections.extend(structured_sections)

    # Build sections now that best_boundaries is known
    for page_num, img_width, line_groups in page_data:
        if not line_groups:
            continue
        for bucket in _section_buckets(line_groups):
            if not bucket:
                continue
            # Per-section boundaries preferred; fall back to document-best
            col_boundaries = _find_col_boundaries(bucket, img_width) or best_boundaries
            rows = _lines_to_rows(bucket, col_boundaries)
            if rows:
                rows = _merge_continuation_rows(rows, col_boundaries)
                title = sanitize_sheet_name(
                    " ".join(c for c in rows[0] if c.strip()) or f"Page {page_num}"
                )
                all_sections.append((title, rows))

    if not all_sections:
        log.info("No text layer found in PDF; attempting OCR fallback.")
        all_sections = _ocr_pdf_pages(content)

    if not all_sections:
        ws = wb.create_sheet(title="Empty PDF")
        ws["A1"] = "No extractable content found in this PDF."
        return save_workbook(wb, stem)

    # structured_sections came from extract_tables() (bordered tables) — don't apply
    # banking-specific _build_single_sheet_rows to them.
    has_structured = bool(structured_sections)
    has_wordbased  = len(all_sections) > len(structured_sections)

    if mode == "single":
        if has_structured and not has_wordbased:
            # Pure grid-extracted PDF — stack tables with a blank-row separator
            combined_rows: list[list[str]] = []
            for i, (_name, rows) in enumerate(all_sections):
                if i > 0:
                    combined_rows.append([])  # blank row separator between tables
                combined_rows.extend(rows)
            ws = wb.create_sheet(title="All Data")
            write_data_to_sheet(ws, combined_rows, "All Data")
        elif has_structured:
            # Mixed: structured tables first (simple concat), then banking sections
            combined_rows = []
            for _name, rows in structured_sections:
                combined_rows.extend(rows)
                combined_rows.append([])
            word_sections = all_sections[len(structured_sections):]
            if word_sections:
                combined_rows.extend(_build_single_sheet_rows(word_sections))
            ws = wb.create_sheet(title="All Data")
            write_data_to_sheet(ws, combined_rows, "All Data")
        else:
            combined_rows = _build_single_sheet_rows(all_sections)
            ws = wb.create_sheet(title="All Data")
            write_data_to_sheet(ws, combined_rows, "All Data")
    else:
        seen_names: dict[str, int] = {}
        for sheet_name, rows in all_sections:
            safe = sanitize_sheet_name(sheet_name)
            if safe in seen_names:
                seen_names[safe] += 1
                safe = sanitize_sheet_name(f"{safe} {seen_names[safe]}")
            else:
                seen_names[safe] = 1
            ws = wb.create_sheet(title=safe)
            write_data_to_sheet(ws, rows, safe)

    return save_workbook(wb, stem)


def convert_docx(content: bytes, stem: str) -> str:
    """DOCX → Excel. Extracts all tables; falls back to paragraph text."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is not installed. Run: pip install python-docx")

    import io

    doc = Document(io.BytesIO(content))
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    sheets_created = 0

    # Extract tables
    for t_idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            sheet_name = f"Table {t_idx}"
            ws = wb.create_sheet(title=sheet_name)
            write_data_to_sheet(ws, rows, sheet_name)
            sheets_created += 1

    # If no tables, extract paragraphs
    if sheets_created == 0:
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            rows = [["#", "Paragraph"]] + [[i, p] for i, p in enumerate(paragraphs, 1)]
            ws = wb.create_sheet(title="Document Text")
            write_data_to_sheet(ws, rows, "Document Text")
            sheets_created += 1

    if sheets_created == 0:
        ws = wb.create_sheet(title="Empty Document")
        ws["A1"] = "No content found in this document."

    return save_workbook(wb, stem)


def _detect_columns_from_gaps(
    line_groups: list[list[dict]], img_width: int
) -> list[int] | None:
    """
    When no header row is detected, infer column boundaries from the distribution
    of large inter-word gaps across all lines.  Gaps that appear consistently in
    the same x-range across multiple lines mark column separators.
    """
    gap_midpoints: list[float] = []

    for line in line_groups:
        if len(line) < 2:
            continue
        words = sorted(line, key=lambda w: w["left"])
        gaps = [
            (words[i + 1]["left"] - words[i]["right"],
             (words[i]["right"] + words[i + 1]["left"]) / 2)
            for i in range(len(words) - 1)
        ]
        sizes = [g for g, _ in gaps]
        if not sizes:
            continue
        median_gap = sorted(sizes)[len(sizes) // 2]
        threshold = max(median_gap * 1.5, 15)
        for gap, mid in gaps:
            if gap >= threshold:
                gap_midpoints.append(mid)

    if not gap_midpoints:
        return None

    gap_midpoints.sort()
    clusters: list[list[float]] = [[gap_midpoints[0]]]
    for mid in gap_midpoints[1:]:
        if mid - clusters[-1][-1] < 40:
            clusters[-1].append(mid)
        else:
            clusters.append([mid])

    # Only trust boundaries that appear in at least 25 % of lines
    min_support = max(2, len(line_groups) // 4)
    boundaries = [0]
    for cluster in clusters:
        if len(cluster) >= min_support:
            boundaries.append(int(sum(cluster) / len(cluster)))
    boundaries.append(img_width)

    return boundaries if len(boundaries) > 2 else None


def _merge_repeating_cols(rows: list[list[str]]) -> list[list[str]]:
    """
    When an image table has repeated column groups (e.g. Name|Phone|Name|Phone|Name|Phone),
    normalize to the base group by stacking each repetition as new rows.

    Detection: find the smallest period P (≥ 2) that divides the column count evenly
    and has at least 2 repetitions.  The period with the most non-empty data wins.
    """
    if not rows:
        return rows

    ncols = max(len(r) for r in rows)
    if ncols < 4:
        return rows

    best_period = None
    best_score  = -1

    for period in range(2, ncols // 2 + 1):
        if ncols % period != 0:
            continue
        reps = ncols // period
        # Score = total non-empty cells when stacked (reward smaller period)
        score = 0
        for row in rows:
            padded = (row + [""] * ncols)[:ncols]
            for r in range(reps):
                chunk = padded[r * period:(r + 1) * period]
                score += sum(1 for c in chunk if c.strip())
        # Prefer the smallest period (penalise larger periods slightly)
        adjusted = score - period * 0.1
        if adjusted > best_score:
            best_score  = adjusted
            best_period = period

    if best_period is None or best_period == ncols:
        return rows

    reps    = ncols // best_period
    stacked = []
    for row in rows:
        padded = (row + [""] * ncols)[:ncols]
        for r in range(reps):
            chunk = padded[r * best_period:(r + 1) * best_period]
            if any(c.strip() for c in chunk):
                stacked.append(chunk)
    return stacked


def _detect_table_grid(gray):
    """
    Detect a drawn table grid in a grayscale numpy image using morphological operations.
    Returns:
      (col_bounds, row_bounds)  — full grid (intersections found)
      (None, row_bounds)        — row-only grid (horizontal lines, no verticals)
      None                      — no grid detected
    """
    import cv2 as _cv2
    import numpy as _np

    h, w = gray.shape

    def _cluster(values, tol=12):
        vals = sorted(set(values))
        if not vals:
            return []
        clusters = [[vals[0]]]
        for v in vals[1:]:
            if v - clusters[-1][-1] <= tol:
                clusters[-1].append(v)
            else:
                clusters.append([v])
        return sorted(sum(c) // len(c) for c in clusters)

    # Global Otsu threshold — better than adaptive for detecting straight lines
    _, binary = _cv2.threshold(gray, 0, 255, _cv2.THRESH_BINARY_INV + _cv2.THRESH_OTSU)

    # Horizontal lines: long thin kernel (at least 1/10 of image width)
    h_kernel_len = max(20, w // 10)
    h_kernel = _cv2.getStructuringElement(_cv2.MORPH_RECT, (h_kernel_len, 1))
    h_lines   = _cv2.morphologyEx(binary, _cv2.MORPH_OPEN, h_kernel)

    # Vertical lines: tall thin kernel (at least 1/10 of image height)
    v_kernel_len = max(20, h // 10)
    v_kernel  = _cv2.getStructuringElement(_cv2.MORPH_RECT, (1, v_kernel_len))
    v_lines   = _cv2.morphologyEx(binary, _cv2.MORPH_OPEN, v_kernel)

    # ── Check for full grid (intersections) ───────────────────────────────────
    intersections = _cv2.bitwise_and(h_lines, v_lines)
    dilate_k = _cv2.getStructuringElement(_cv2.MORPH_RECT, (7, 7))
    intersections = _cv2.dilate(intersections, dilate_k)
    contours, _ = _cv2.findContours(intersections, _cv2.RETR_EXTERNAL, _cv2.CHAIN_APPROX_SIMPLE)

    if contours:
        centres = []
        for cnt in contours:
            x, y, cw, ch = _cv2.boundingRect(cnt)
            centres.append((x + cw // 2, y + ch // 2))
        xs = [c[0] for c in centres]
        ys = [c[1] for c in centres]
        col_bounds = _cluster(xs, tol=12)
        row_bounds = _cluster(ys, tol=12)
        if len(col_bounds) >= 3 and len(row_bounds) >= 3:
            return col_bounds, row_bounds

    return None


def _detect_header_band(gray):
    """
    Detect a solid coloured header band (e.g. purple transaction-table header
    with white text).  Returns (band_top, band_bottom) in pixel rows, or None.

    Works on any background tint by using RELATIVE darkness: the header band is
    markedly darker than the page-background median luminance.  This handles
    both white-background documents and dark/tinted screenshots.
    """
    import numpy as _np
    h, w = gray.shape
    cx0, cx1 = w // 5, 4 * w // 5
    row_lum    = gray[:, cx0:cx1].mean(axis=1)        # mean luminance per row
    median_lum = float(_np.median(row_lum))           # page-background level
    dark_thresh = median_lum * 0.65                   # band is ≥35 % darker
    is_dark = row_lum < dark_thresh

    # Collect runs of consecutive dark rows
    bands = []
    r = 0
    while r < h:
        if is_dark[r]:
            start = r
            while r < h and is_dark[r]:
                r += 1
            bands.append((start, r))
        else:
            r += 1

    # Return the first band of plausible header height
    max_h = max(60, h // 10)
    for start, end in bands:
        if 8 <= (end - start) <= max_h:
            return start, end
    return None


def _columns_from_header_band(pil_img, band_top, band_bottom):
    """
    OCR a coloured header band by inverting it (white-text-on-dark → dark-on-
    light) and return column boundary x-positions derived from the header
    word positions.  Returns a sorted boundary list [0, b1, b2, ..., width]
    or None if fewer than 3 header words are found.
    """
    try:
        import pytesseract
        from pytesseract import Output
        from PIL import Image as _PILImage
    except ImportError:
        return None
    import cv2 as _cv2
    import numpy as _np

    pad = 4
    top = max(0, band_top - pad)
    bot = min(pil_img.height, band_bottom + pad)
    band = pil_img.crop((0, top, pil_img.width, bot))

    gray = _np.array(band.convert("L"))
    # Upscale for better OCR of the (often small) header text
    scale = 3
    gray = _cv2.resize(gray, (gray.shape[1] * scale, gray.shape[0] * scale),
                       interpolation=_cv2.INTER_LANCZOS4)
    # Invert: white text on dark BG → black text on light BG
    inv = 255 - gray
    _, bw = _cv2.threshold(inv, 0, 255, _cv2.THRESH_BINARY + _cv2.THRESH_OTSU)

    data = pytesseract.image_to_data(
        _PILImage.fromarray(bw), output_type=Output.DICT, config="--oem 3 --psm 6"
    )
    words = []
    for i, txt in enumerate(data["text"]):
        t = txt.strip()
        if t and int(data["conf"][i]) > 20:
            words.append({
                "text": t,
                "left":  data["left"][i] // scale,
                "right": (data["left"][i] + data["width"][i]) // scale,
            })
    if len(words) < 3:
        return None

    words.sort(key=lambda x: x["left"])
    # Column boundaries: midpoint between the right edge of word i and the
    # left edge of word i+1.
    bounds = [0]
    for i in range(len(words) - 1):
        bounds.append((words[i]["right"] + words[i + 1]["left"]) // 2)
    bounds.append(pil_img.width)

    # Clean header labels: strip non-ASCII OCR noise, then snap to the closest
    # known column name when there's a clear match (fixes e.g. "O�SCRIPTION").
    _KNOWN = ["DATE", "DESCRIPTION", "WITHDRAWAL", "DEPOSIT", "BALANCE",
              "CREDIT", "DEBIT", "AMOUNT", "REFERENCE", "TYPE"]

    import difflib

    def _clean_label(raw: str) -> str:
        cleaned = "".join(c for c in raw if c.isalnum() or c in " $.,-/").strip()
        up = cleaned.upper()
        # Snap to the closest known header via sequence similarity
        best, best_score = cleaned, 0.0
        for k in _KNOWN:
            score = difflib.SequenceMatcher(None, up, k).ratio()
            if score > best_score:
                best, best_score = k, score
        return best if best_score >= 0.6 else cleaned

    labels = [_clean_label(w_["text"]) for w_ in words]
    return bounds, labels


def _extract_cells_from_grid(pil_img, col_bounds, row_bounds):
    """
    Given detected grid line positions, crop each cell from pil_img and OCR it
    with PSM 7 (single text line) — far more accurate than whole-image OCR for small cells.
    Returns list-of-rows (list[list[str]]), empty rows filtered out.
    """
    try:
        import pytesseract
        from pytesseract import Output
        from PIL import Image as _PILImage
    except ImportError:
        return []

    import cv2 as _cv2
    import numpy as _np

    # Scale factor: work on a 2x upscaled image for better per-cell OCR
    scale = 2
    img_big = pil_img.resize((pil_img.width * scale, pil_img.height * scale),
                              resample=_PILImage.LANCZOS)
    gray_big = _np.array(img_big.convert("L"))

    cfg = "--oem 3 --psm 7"
    rows: list[list[str]] = []

    for ri in range(len(row_bounds) - 1):
        y0 = max(0, row_bounds[ri]  * scale - 2)
        y1 = min(gray_big.shape[0], row_bounds[ri + 1] * scale + 2)
        row_cells: list[str] = []
        for ci in range(len(col_bounds) - 1):
            x0 = max(0, col_bounds[ci]  * scale - 2)
            x1 = min(gray_big.shape[1], col_bounds[ci + 1] * scale + 2)
            cell_gray = gray_big[y0:y1, x0:x1]
            if cell_gray.size == 0:
                row_cells.append("")
                continue
            # Adaptive threshold for the cell crop
            _, cell_bw = _cv2.threshold(cell_gray, 0, 255,
                                        _cv2.THRESH_BINARY + _cv2.THRESH_OTSU)
            cell_img = _PILImage.fromarray(cell_bw)
            text = pytesseract.image_to_string(cell_img, config=cfg).strip()
            row_cells.append(text)
        if any(c for c in row_cells):
            rows.append(row_cells)

    return rows


_SCORE_GARBAGE_RE = _re.compile(r'^[\W_]+$')          # only punctuation/symbols
_SCORE_REAL_RE    = _re.compile(r'[A-Za-z0-9]{2,}')   # a real token (≥2 alnum)

def _score_table_rows(rows: list[list[str]]) -> float:
    """
    Rate how 'table-like' a candidate result is (higher = better).  Pure function
    over the rows — no image needed.  Used to pick the best of several extraction
    strategies (grid / header-band / word-based) so a strategy that mis-fires on a
    given image simply loses to a cleaner candidate.

    Signals:
      + column-count consistency (most rows share the same #cells)
      + content quality (cells with real ≥2-char tokens)
      + fill ratio (few empty cells)
      - garbage cells (lone punctuation, the U+FFFD replacement char, 1-char noise)
      - implausible shape (too many columns, mostly-garbage)
      * scaled by data-row count so a clean multi-row table beats a tiny fluke
    """
    if not rows:
        return 0.0

    # Drop fully-empty rows for measurement
    data_rows = [r for r in rows if any(str(c).strip() for c in r)]
    if not data_rows:
        return 0.0

    n_rows = len(data_rows)
    col_counts = [len(r) for r in data_rows]
    max_cols   = max(col_counts)

    # Cell-level tallies
    total_cells = 0
    filled      = 0
    real        = 0
    garbage     = 0
    for r in data_rows:
        for c in r:
            total_cells += 1
            s = str(c).strip()
            if not s:
                continue
            filled += 1
            if "�" in s or (len(s) == 1 and not s.isalnum()) or _SCORE_GARBAGE_RE.match(s):
                garbage += 1
            elif _SCORE_REAL_RE.search(s):
                real += 1

    if filled == 0:
        return 0.0

    fill_ratio    = filled / total_cells
    real_ratio    = real / filled
    garbage_ratio = garbage / filled
    real_per_row  = real / n_rows      # real cells per row = genuine column structure

    # Column-count consistency: fraction of rows sharing the modal column count
    from collections import Counter
    modal_count, modal_freq = Counter(col_counts).most_common(1)[0]
    consistency = modal_freq / n_rows

    score = 0.0
    score += consistency   * 2.0     # regular shape
    score += real_ratio    * 2.5     # cells must hold REAL tokens, not fragments
    score += fill_ratio    * 0.5     # density
    score += min(real_per_row, 8.0) * 0.6   # reward real content split ACROSS columns
    score -= garbage_ratio * 5.0     # punish OCR noise hard

    # Sanity penalties for implausible shapes.  Real business tables rarely
    # exceed ~8 columns; a bogus over-segmented grid (e.g. 16 cols of OCR
    # fragments) is heavily penalised here so it loses to a clean candidate.
    if max_cols > 8:
        score -= (max_cols - 8) * 0.7
    if garbage_ratio > 0.5:
        score -= 3.0

    # Scale by data volume (saturating) so clean multi-row tables win
    import math
    score *= (1.0 + math.log1p(n_rows) * 0.15)

    return score


def _columns_from_left_edges(line_groups, img_width):
    """
    Detect columns by clustering word LEFT edges across all lines.  Works well
    for left-aligned tables (e.g. employee/project tables) where each column
    starts at a consistent x even when row content length varies.  Returns a
    boundary list [0, b1, ..., img_width] or None.
    """
    lefts = sorted(w["left"] for ln in line_groups for w in ln)
    if len(lefts) < 4:
        return None
    tol = max(20, img_width // 26)
    clusters = [[lefts[0]]]
    for x in lefts[1:]:
        if x - clusters[-1][-1] <= tol:
            clusters[-1].append(x)
        else:
            clusters.append([x])
    # Keep clusters that appear in a meaningful share of lines (real columns)
    n = len(line_groups)
    sig = [c for c in clusters if len(c) >= max(2, n * 0.4)]
    if len(sig) < 2:
        return None
    centres = sorted(sum(c) // len(c) for c in sig)
    bounds = [0]
    for i in range(len(centres) - 1):
        bounds.append((centres[i] + centres[i + 1]) // 2)
    bounds.append(img_width)
    return bounds if len(bounds) > 2 else None


def _columns_from_ruled_lines(gray, img_width):
    """
    Detect FAINT vertical table borders on a light-background ruled table (e.g.
    SBI/bank statement screenshots with thin grey gridlines) and return column
    x-boundaries.  Otsu misses these lines because they're only slightly darker
    than the near-white page, so use a fixed high threshold; long-kernel
    morphology isolates the vertical rules from text.  Returns a boundary list
    [0, b1, ..., img_width] or None.

    Gated to light backgrounds (mean luminance > 200) so it can't fabricate a
    grid from the foreground of a dark/inverted image.
    """
    import cv2 as _cv2
    import numpy as _np

    h, w = gray.shape
    if float(gray.mean()) <= 200:
        return None

    binary = (gray < 240).astype("uint8") * 255          # catch faint grey rules
    v_kernel = _cv2.getStructuringElement(_cv2.MORPH_RECT, (1, max(20, h // 8)))
    v_lines = _cv2.morphologyEx(binary, _cv2.MORPH_OPEN, v_kernel)
    contours, _ = _cv2.findContours(v_lines, _cv2.RETR_EXTERNAL, _cv2.CHAIN_APPROX_SIMPLE)

    xs = []
    for c in contours:
        x, y, cw, ch = _cv2.boundingRect(c)
        if ch > h * 0.2:                                 # real rules span much of the height
            xs.append(x + cw // 2)
    if len(xs) < 3:
        return None

    xs.sort()
    clusters = [[xs[0]]]
    for x in xs[1:]:
        if x - clusters[-1][-1] <= 15:
            clusters[-1].append(x)
        else:
            clusters.append([x])
    centres = [sum(g) // len(g) for g in clusters]

    # Build boundaries, then merge columns narrower than a min width — this drops
    # the thin empty columns created by the table's outer border lines.
    bounds = sorted(set([0] + centres + [w]))
    min_w = max(15, img_width // 25)
    merged = [bounds[0]]
    for b in bounds[1:]:
        if b - merged[-1] < min_w:
            continue
        merged.append(b)
    if merged[-1] < w:
        merged[-1] = w
    return merged if len(merged) >= 4 else None          # need ≥3 real columns


def _header_row_for_ruled(pil_img, gray, col_bounds):
    """
    Find the column-header row of a ruled table (the full-page OCR pass often
    skips it because it sits sandwiched between horizontal rules) and OCR it on
    its own.  Returns (header_cells, band_top) mapped to col_bounds, or None.

    Strategy: detect horizontal rules (fixed threshold, like the vertical ones),
    OCR each inter-rule band top-to-bottom, and take the first band that contains
    ≥4 distinct table-header keywords.
    """
    try:
        import pytesseract
        from pytesseract import Output
        from PIL import Image as _PILImage
    except ImportError:
        return None
    import cv2 as _cv2
    import numpy as _np

    h, w = gray.shape
    binary = (gray < 240).astype("uint8") * 255
    h_kernel = _cv2.getStructuringElement(_cv2.MORPH_RECT, (max(20, w // 8), 1))
    h_lines = _cv2.morphologyEx(binary, _cv2.MORPH_OPEN, h_kernel)
    ys = [y for y in range(h) if h_lines[y].sum() / 255 > w * 0.4]
    if len(ys) < 2:
        return None
    clusters = [[ys[0]]]
    for y in ys[1:]:
        if y - clusters[-1][-1] <= 8:
            clusters[-1].append(y)
        else:
            clusters.append([y])
    rules = [sum(g) // len(g) for g in clusters]

    _HDR_KW = {"date", "narration", "description", "ref", "cheque", "debit", "credit",
               "balance", "withdrawal", "deposit", "particulars", "value", "amount",
               "transaction", "type", "no", "chq", "dr", "cr"}
    for i in range(len(rules) - 1):
        t, b = rules[i], rules[i + 1]
        if not (12 <= b - t <= max(60, h // 8)):
            continue
        crop = pil_img.crop((0, max(0, t - 1), w, b + 1))
        cg = _np.array(crop.convert("L"))
        cg = _cv2.resize(cg, (cg.shape[1] * 3, cg.shape[0] * 3),
                         interpolation=_cv2.INTER_LANCZOS4)
        _, cb = _cv2.threshold(cg, 0, 255, _cv2.THRESH_BINARY + _cv2.THRESH_OTSU)
        d = pytesseract.image_to_data(
            _PILImage.fromarray(cb), output_type=Output.DICT, config="--oem 3 --psm 6"
        )
        words = [
            {"text": d["text"][j].strip(),
             "left": d["left"][j] // 3,
             "right": (d["left"][j] + d["width"][j]) // 3}
            for j in range(len(d["text"]))
            if d["text"][j].strip() and int(d["conf"][j]) >= 0
        ]
        hits = sum(1 for wd in words if wd["text"].lower().strip(".():|") in _HDR_KW)
        if hits >= 4:
            cells = _line_to_cells_fixed(words, col_bounds)
            cells = [c.strip().strip("|").strip() for c in cells]   # drop border pipes
            return cells, t
    return None


# Transaction keywords that OCR sometimes glues to the preceding word in
# all-caps statements (e.g. "POSPURCHASE", "PREAUTHORIZEDCREDIT").
_GLUE_KEYWORDS = ("PURCHASE", "CREDIT", "DEBIT", "WITHDRAWAL", "DEPOSIT",
                  "CHARGE", "PAYMENT", "TRANSFER", "INTEREST", "BALANCE")
_GLUE_RE = _re.compile(r"^([A-Z]{2,})(" + "|".join(_GLUE_KEYWORDS) + r")$")

def _desplit_glued(text: str) -> str:
    """
    Re-insert a missing space when OCR glued a known transaction keyword onto the
    end of an all-caps word: 'POSPURCHASE' → 'POS PURCHASE', 'ATMWITHDRAWAL' →
    'ATM WITHDRAWAL'.  Deliberately conservative: only all-caps tokens of the exact
    form <prefix><keyword> where the keyword is at the END, so it can't damage
    normal words (e.g. 'ACCREDITED', 'CREDITOR' are left untouched).
    """
    m = _GLUE_RE.match(text)
    if m:
        return f"{m.group(1)} {m.group(2)}"
    return text


def _azure_tables_to_rows(table) -> list[list[str]]:
    """Convert one Azure Document Intelligence table object to a list-of-rows grid."""
    n_rows = int(getattr(table, "row_count", 0) or 0)
    n_cols = int(getattr(table, "column_count", 0) or 0)
    if n_rows <= 0 or n_cols <= 0:
        return []
    grid = [["" for _ in range(n_cols)] for _ in range(n_rows)]
    for cell in (table.cells or []):
        r = int(getattr(cell, "row_index", 0) or 0)
        c = int(getattr(cell, "column_index", 0) or 0)
        if 0 <= r < n_rows and 0 <= c < n_cols:
            grid[r][c] = (getattr(cell, "content", "") or "").strip()
    return [row for row in grid if any(v.strip() for v in row)]


def convert_image_azure(content: bytes, stem: str, mode: str = "single",
                        merge_cols: bool = False) -> str:
    """
    Image → Excel via Azure AI Document Intelligence (prebuilt-layout), which
    returns native table structure + high-accuracy OCR.  Raises CloudOCRUnavailable
    when the SDK isn't installed, credentials are missing, or the API call fails,
    so the caller can fall back to the Tesseract pipeline.
    """
    if not _azure_configured():
        raise CloudOCRUnavailable("Azure Document Intelligence credentials are not configured.")
    try:
        from azure.core.credentials import AzureKeyCredential
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
    except ImportError as e:
        raise CloudOCRUnavailable(
            "azure-ai-documentintelligence is not installed. "
            "Run: pip install azure-ai-documentintelligence"
        ) from e

    try:
        client = DocumentIntelligenceClient(
            endpoint=AZURE_DI_ENDPOINT,
            credential=AzureKeyCredential(AZURE_DI_KEY),
        )
        # bytes_source is the version-stable way to send local image bytes
        # (the SDK base64-encodes it); the raw-bytes `body=` form is flakier.
        poller = client.begin_analyze_document(
            "prebuilt-layout",
            AnalyzeDocumentRequest(bytes_source=content),
        )
        result = poller.result()
    except Exception as e:
        raise CloudOCRUnavailable(f"Azure Document Intelligence call failed: {e}") from e

    # Build (title, rows) sections from detected tables.
    sections: list[tuple[str, list[list[str]]]] = []
    for idx, table in enumerate(getattr(result, "tables", None) or [], start=1):
        rows = _azure_tables_to_rows(table)
        if merge_cols:
            rows = _merge_repeating_cols(rows)
        if rows:
            title = sanitize_sheet_name(
                " ".join(c for c in rows[0] if c.strip()) or f"Table {idx}"
            )
            sections.append((title, rows))

    # No tables detected — fall back to the document's plain text content.
    fallback_text = (getattr(result, "content", "") or "").strip()
    return _cloud_sections_to_workbook(sections, stem, mode, fallback_text)


def _cloud_sections_to_workbook(sections, stem, mode, fallback_text=""):
    """Shared writer for cloud-OCR engines: list of (title, rows) → .xlsx.
    single mode stacks tables with a blank-row separator; separate mode gives one
    sheet per table (de-duplicating sheet names)."""
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    if sections:
        if mode == "single":
            combined: list[list[str]] = []
            for i, (_name, rows) in enumerate(sections):
                if i > 0:
                    combined.append([])          # blank-row separator between tables
                combined.extend(rows)
            ws = wb.create_sheet(title="All Data")
            write_data_to_sheet(ws, combined, "All Data")
        else:
            seen: dict[str, int] = {}
            for name, rows in sections:
                safe = sanitize_sheet_name(name)
                if safe in seen:
                    seen[safe] += 1
                    safe = sanitize_sheet_name(f"{safe} {seen[safe]}")
                else:
                    seen[safe] = 1
                ws = wb.create_sheet(title=safe)
                write_data_to_sheet(ws, rows, safe)
    else:
        rows = [[line] for line in fallback_text.splitlines() if line.strip()] or [["No content found."]]
        ws = wb.create_sheet(title="Extracted Text")
        write_data_to_sheet(ws, rows, "Extracted Text")

    return save_workbook(wb, stem)


# Override via env GEMINI_MODEL if this default isn't available on your free tier.
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash").strip() or "gemini-2.5-flash"
# Used automatically when the primary model returns 429 (quota/rate-limit exceeded).
GEMINI_FALLBACK_MODEL = os.getenv("GEMINI_FALLBACK_MODEL", "gemini-2.0-flash").strip()


def _parse_gemini_retry_delay(err) -> Optional[float]:
    """Pull the API-suggested retry delay (seconds) from a Gemini error, if present.
    Matches RetryInfo (`'retryDelay': '33s'`) or the human text (`retry in 33.8s`)."""
    s = str(err)
    m = _re.search(r"retry(?:Delay)?['\":\s]+(?:in\s+)?(\d+(?:\.\d+)?)\s*s", s, _re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            return None
    return None


def _sniff_image_mime(data: bytes) -> str:
    if data[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    return "image/png"


_GEMINI_PROMPT = (
    "You are an OCR table extractor. Extract ALL data from this image as tables. "
    "Return ONLY JSON of the form "
    '{"tables": [{"rows": [["cell", "cell"], ["cell", "cell"]]}]}. '
    "Rules: preserve every row and column; use an empty string for blank cells; "
    "keep numbers, dates, and text exactly as written (do not reformat); include the "
    "header row. If a region is key-value text rather than a grid, put each line as a "
    "single-cell row. Output no commentary, only the JSON."
)


def convert_image_gemini(content: bytes, stem: str, mode: str = "single",
                         merge_cols: bool = False) -> str:
    """
    Image → Excel via Google Gemini (AI Studio, free tier — no credit card).
    Gemini reads the image and returns table data as JSON, which we map to sheets.
    Raises CloudOCRUnavailable when the SDK isn't installed, the key is missing, or
    the API call fails, so the caller can fall back to Tesseract.
    """
    if not _gemini_configured():
        raise CloudOCRUnavailable("Gemini API key (GEMINI_API_KEY) is not configured.")
    try:
        from google import genai
        from google.genai import types
    except ImportError as e:
        raise CloudOCRUnavailable(
            "google-genai is not installed. Run: pip install google-genai"
        ) from e

    import time as _time
    client = genai.Client(api_key=GEMINI_API_KEY)
    contents = [
        types.Part.from_bytes(data=content, mime_type=_sniff_image_mime(content)),
        _GEMINI_PROMPT,
    ]
    cfg = types.GenerateContentConfig(response_mime_type="application/json")

    # Retry policy (per model), then fall back to the next model on HARD exhaustion:
    #   • per-minute rate limit (RPM/TPM) or 503/overloaded → transient: wait the
    #     API-suggested delay (capped) and retry the SAME model. Free-tier 2.5-flash
    #     is 5 RPM, so bursts trip this — but it clears within the minute.
    #   • per-day quota (RPD) or "limit: 0" (model not on this tier) → this model is
    #     done → try the next model.
    #   • anything else (400/401/403) → permanent → abort (next model fails too).
    models = []
    for m in (GEMINI_MODEL, GEMINI_FALLBACK_MODEL):
        if m and m not in models:
            models.append(m)

    _WAIT_CAP = 30.0          # never sleep longer than this per wait
    _MAX_ATTEMPTS = 3         # up to 2 waited retries per model

    raw = None
    last_err = None
    for model in models:
        for attempt in range(_MAX_ATTEMPTS):
            try:
                response = client.models.generate_content(
                    model=model, contents=contents, config=cfg
                )
                raw = (response.text or "").strip()
                log.info("Gemini OCR used model %s", model)
                break
            except Exception as e:
                last_err = e
                msg = str(e).lower()
                is_429       = "429" in msg or "resource_exhausted" in msg or "rate limit" in msg
                is_503       = any(t in msg for t in ("503", "500", "unavailable", "overloaded", "high demand"))
                is_limit0    = "limit: 0" in msg or "limit:0" in msg
                is_per_day   = "perday" in msg or "per day" in msg
                is_per_min   = "perminute" in msg or "per minute" in msg
                retry_delay  = _parse_gemini_retry_delay(e)

                # Hard exhaustion for this model → move on to the next model.
                if is_limit0 or (is_429 and is_per_day):
                    log.warning("Gemini model %s hard-exhausted (daily/limit:0); trying next model.", model)
                    break

                # Transient: per-minute 429, a 429 with a suggested delay, or 503.
                if is_503 or (is_429 and (is_per_min or retry_delay is not None)):
                    if attempt < _MAX_ATTEMPTS - 1:
                        wait = min(retry_delay if retry_delay is not None else 1.5 * (attempt + 1), _WAIT_CAP)
                        log.warning("Gemini model %s transient (%s); waiting %.1fs and retrying.",
                                    model, "429/rate-limit" if is_429 else "503", wait)
                        _time.sleep(wait)
                        continue
                    break                                   # retries exhausted → next model

                # Unknown 429 (no day/min signal, no delay) → treat as this-model-done.
                if is_429:
                    log.warning("Gemini model %s rate-limited (unclassified); trying next model.", model)
                    break

                raise CloudOCRUnavailable(f"Gemini API call failed: {e}") from e
        if raw is not None:
            break
    if raw is None:
        raise CloudOCRUnavailable(f"Gemini API call failed (all models exhausted): {last_err}")

    # Parse the JSON (tolerate stray code fences if the model added them).
    import json
    if raw.startswith("```"):
        raw = raw.strip("`")
        raw = raw[raw.find("{"):] if "{" in raw else raw
    try:
        data = json.loads(raw)
    except Exception as e:
        raise CloudOCRUnavailable(f"Gemini returned unparseable JSON: {e}") from e

    sections: list[tuple[str, list[list[str]]]] = []
    for idx, table in enumerate(data.get("tables", []) or [], start=1):
        rows = [
            [("" if c is None else str(c)).strip() for c in (row or [])]
            for row in (table.get("rows", []) or [])
        ]
        rows = [r for r in rows if any(c for c in r)]
        if merge_cols:
            rows = _merge_repeating_cols(rows)
        if rows:
            title = sanitize_sheet_name(
                " ".join(c for c in rows[0] if c.strip()) or f"Table {idx}"
            )
            sections.append((title, rows))

    return _cloud_sections_to_workbook(sections, stem, mode, fallback_text=raw)


_RAPIDOCR_ENGINE = None   # lazily-initialised singleton (ONNX model load is slow)

def _ocr_words_rapidocr(pil_img):
    """
    OCR an image with RapidOCR (a free, local PP-OCR neural model via ONNXRuntime —
    no card, no quota, no internet).  Returns words in the same shape convert_image's
    `all_words` uses, or None if RapidOCR isn't installed / finds nothing.  Used as a
    higher-accuracy fallback for low-confidence (degraded photo) images.
    """
    global _RAPIDOCR_ENGINE
    try:
        import numpy as _np
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        return None
    try:
        if _RAPIDOCR_ENGINE is None:
            _RAPIDOCR_ENGINE = RapidOCR()
        result, _ = _RAPIDOCR_ENGINE(_np.array(pil_img.convert("RGB")))
    except Exception as e:
        log.warning("RapidOCR failed: %s", e)
        return None
    if not result:
        return None

    words = []
    for item in result:
        box, text = item[0], item[1]
        t = (text or "").strip()
        if not t:
            continue
        xs = [p[0] for p in box]
        ys = [p[1] for p in box]
        left, right = int(min(xs)), int(max(xs))
        top, bottom = int(min(ys)), int(max(ys))
        words.append({
            "text":   _desplit_glued(t),
            "left":   left,
            "top":    top,
            "right":  right,
            "width":  max(right - left, 1),
            "height": max(bottom - top, 1),
        })
    return words or None


# ── Spreadsheet-screenshot chrome cleanup ──────────────────────────────────────
# Photos of an Excel/Sheets window also capture the app's own UI: the Name Box
# ("A26"), the column-letter band ("A B C D …"), and the row-number gutter
# ("2", "3", …).  OCR folds these into the data — the column-letter band lands as
# a spurious first row that shoves the first contacts' phones down a row, and the
# gutter numbers glue onto names ("2 viswa").  These are *not* real data, so we
# strip them in a final validation pass — but only when we're confident the image
# is a spreadsheet screenshot, so ordinary photos/scans are never touched.
_COL_LETTER_CELL_RE = _re.compile(r"^[A-Z]( [A-Z])*$")   # "B", "C D E", "F"
_NAME_BOX_RE        = _re.compile(r"^[A-Z]{1,3}\d{1,7}$")  # "A26", "AB12"
_ROW_GUTTER_RE      = _re.compile(r"^(?:\d{1,3}[\s.]+)+(?=\D)")  # leading "2 " / "3 2 " before text
_PURE_ROWNUM_RE     = _re.compile(r"^\d{1,3}$")            # a lone row number
_TRAIL_COL_LETTER_RE = _re.compile(r"\s+[A-Z]$")          # name with " H" stuck on

# Only the top few rows carry the header band — bound chrome edits to them so a
# legitimate single-letter or leading-number cell deeper in the table is safe.
_CHROME_TOP_ROWS = 3


def _looks_like_spreadsheet_screenshot(rows: list[list[str]]) -> bool:
    """True if the OCR rows carry Excel/Sheets UI chrome (name box / letter band)."""
    col_letters = 0
    for row in rows[:_CHROME_TOP_ROWS]:
        for cell in row:
            s = cell.strip()
            if _COL_LETTER_CELL_RE.match(s):
                col_letters += s.count(" ") + 1   # "C D E" → 3 column letters
    name_box = False
    for row in rows[:2]:
        nonempty = [c.strip() for c in row if c.strip()]
        if len(nonempty) == 1 and _NAME_BOX_RE.match(nonempty[0]):
            name_box = True
    return name_box or col_letters >= 2


def _strip_spreadsheet_chrome(rows: list[list[str]]) -> list[list[str]]:
    """Remove Excel UI chrome (name box, column-letter band, row-number gutter).

    No-op unless the rows look like a spreadsheet screenshot.  Returns a new,
    chrome-free row list (rows emptied by the cleanup are dropped).
    """
    if not _looks_like_spreadsheet_screenshot(rows):
        return rows

    cleaned: list[list[str]] = []
    for idx, row in enumerate(rows):
        # Drop the Name Box artifact: a lone "A26"-style token, rest of row empty.
        nonempty = [c.strip() for c in row if c.strip()]
        if len(nonempty) == 1 and _NAME_BOX_RE.match(nonempty[0]):
            continue

        new = list(row)
        for ci, cell in enumerate(new):
            s = cell.strip()
            # Pure column-letter cells in the header band → blank.
            if idx < _CHROME_TOP_ROWS and _COL_LETTER_CELL_RE.match(s):
                new[ci] = ""
                continue
            # Row-number gutter lives in the leftmost column only.
            if ci == 0:
                if _PURE_ROWNUM_RE.match(s):
                    new[ci] = ""
                    continue
                stripped = _ROW_GUTTER_RE.sub("", s)
                if stripped != s:
                    new[ci] = stripped.strip()
            # A column letter glued onto a name in the header band ("shaki H").
            if idx < _CHROME_TOP_ROWS and _TRAIL_COL_LETTER_RE.search(new[ci].strip()):
                new[ci] = _TRAIL_COL_LETTER_RE.sub("", new[ci].strip())

        if any(c.strip() for c in new):
            cleaned.append(new)

    dropped = len(rows) - len(cleaned)
    log.info("Spreadsheet chrome detected: stripped UI artifacts (%d row(s) removed).", dropped)
    return cleaned


# A cell that glued a name onto its phone number(s): "pargunan 7981233678",
# "shyam 8978973222", "ramana reddy 8848436887".  The name part must contain a
# letter (so a pure two-phone cell like "9980831997 9963393260" is NOT a match),
# and the trailing part is one or more long digit runs (phone numbers).
_NAME_PHONE_RE = _re.compile(r"^(.+?)\s+(\d{9,}(?:\s+\d{9,})*)$")


def _name_phone_split(cell: str):
    """Return (name, phones) if the cell is 'name + phone number(s)', else None."""
    m = _NAME_PHONE_RE.match(cell.strip())
    if not m:
        return None
    name = m.group(1).strip()
    if not _re.search(r"[A-Za-z]", name):   # the left part must look like a name
        return None
    return name, m.group(2).strip()


def _split_glued_name_phone_columns(rows: list[list[str]]) -> list[list[str]]:
    """Split a column that glued names onto phone numbers into two columns.

    When OCR misses the gap between a name column and its phone column, every cell
    reads as "name 9876543210".  If a column is *predominantly* such cells, split
    it into a name column followed by a phone column.  Other columns (pure names,
    pure phones, descriptions) never match the majority test and pass through.
    """
    if not rows:
        return rows
    ncols = max((len(r) for r in rows), default=0)
    grid = [list(r) + [""] * (ncols - len(r)) for r in rows]

    new_cols: list[list[str]] = []
    split_any = False
    for ci in range(ncols):
        col = [grid[ri][ci].strip() for ri in range(len(grid))]
        nonempty = [c for c in col if c]
        matches = sum(1 for c in nonempty if _name_phone_split(c))
        # Split only when the column is clearly a glued name+phone column: a strong
        # majority of its filled cells carry a trailing phone number.
        if nonempty and matches >= 3 and matches >= 0.6 * len(nonempty):
            split_any = True
            names, phones = [], []
            for c in col:
                parts = _name_phone_split(c)
                if parts:
                    names.append(parts[0])
                    phones.append(parts[1])
                else:
                    names.append(c)
                    phones.append("")
            new_cols.append(names)
            new_cols.append(phones)
        else:
            new_cols.append(col)

    if not split_any:
        return rows

    # Drop any fully-empty trailing columns the split may have exposed.
    while new_cols and not any(c.strip() for c in new_cols[-1]):
        new_cols.pop()

    out = [[new_cols[ci][ri] for ci in range(len(new_cols))] for ri in range(len(grid))]
    log.info("Split a glued name+phone column into separate name/phone columns.")
    return out


# Two long numbers glued into one cell ("9980831997 9963393260") — the signature
# of a row that absorbed a neighbour's phone because a tilted photo smeared the
# y-coordinate row grouping. Used to detect (and, after re-grouping, fix) that.
_TWO_LONG_NUMS_RE = _re.compile(r"\d{9,}\s+\d{9,}")


def _count_number_collisions(rows: list[list[str]]) -> int:
    return sum(1 for r in rows for c in r if _TWO_LONG_NUMS_RE.search(c.strip()))


def _group_lines_tilted(words: list[dict], slope: float, y_tol: float) -> list[list[dict]]:
    """Y-band line grouping (same as the main path) but on a tilt-corrected Y:
    `cy - left*slope`. slope=0 reproduces the untilted grouping exactly."""
    def adj(w):
        return w["top"] + w["height"] / 2 - w["left"] * slope

    sw = sorted(words, key=lambda w: (adj(w), w["left"]))
    lines: list[list[dict]] = []
    cur: list[dict] = [sw[0]]
    cur_a = adj(sw[0])
    for wd in sw[1:]:
        a = adj(wd)
        if abs(a - cur_a) <= y_tol:
            cur.append(wd)
            cur_a = sum(adj(w) for w in cur) / len(cur)
        else:
            lines.append(sorted(cur, key=lambda w: w["left"]))
            cur, cur_a = [wd], a
    lines.append(sorted(cur, key=lambda w: w["left"]))
    return lines


def _refine_tilted_rows(rows: list[list[str]], all_words: list[dict],
                        img_width: int, y_tol: float) -> list[list[str]]:
    """Fix row mis-grouping caused by a tilted photo.

    A tilt smears the y-band grouping so a row absorbs the next row's phone number
    (two numbers land in one cell).  Only when that defect is present, re-group the
    words at a few tilt angles, re-detect columns, and keep the version that has the
    FEWEST such collisions (tie-break: higher table score).  Because it triggers
    only on the defect and must strictly reduce it, straight images are never
    touched.  Returns the original rows if nothing beats them.
    """
    base = _count_number_collisions(rows)
    if base == 0 or len(all_words) < 12:
        return rows

    best_rows = rows
    best_key = (base, -_score_table_rows(rows))
    for slope in (0.02, 0.035, 0.05):
        lgs = _group_lines_tilted(all_words, slope, y_tol)
        for bounds in (_columns_from_left_edges(lgs, img_width),
                       _detect_columns_from_gaps(lgs, img_width)):
            if not bounds or len(bounds) <= 2:
                continue
            cand = []
            for ln in lgs:
                cells = _line_to_cells_fixed(ln, bounds)
                if any(c.strip() for c in cells):
                    cand.append(cells)
            key = (_count_number_collisions(cand), -_score_table_rows(cand))
            if key < best_key:
                best_key, best_rows = key, cand

    if best_rows is not rows:
        log.info("Tilt refinement: number-collisions %d → %d (re-grouped a tilted photo).",
                 base, best_key[0])
    return best_rows


def convert_image(content: bytes, stem: str, filename: str, merge_cols: bool = False) -> str:
    """Image (PNG/JPG) → Excel via OCR (pytesseract, or RapidOCR for hard photos)."""
    try:
        import pytesseract
        from PIL import Image
        from pytesseract import Output
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=(
                "pytesseract or Pillow is not installed. "
                "Run: pip install pytesseract Pillow  "
                "Also install Tesseract OCR engine from https://github.com/UB-Mannheim/tesseract/wiki"
            ),
        )

    import io, os, shutil

    # Auto-locate Tesseract on Windows when it is not in the process PATH
    if not shutil.which("tesseract"):
        _win_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.expandvars(r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"),
        ]
        for _p in _win_paths:
            if os.path.isfile(_p):
                pytesseract.pytesseract.tesseract_cmd = _p
                break

    try:
        img = Image.open(io.BytesIO(content))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        # ── Grid detection (jpgtoexcel.com / AWS Textract technique) ─────────
        # Detect drawn table lines first.  If a grid is found, OCR each cell
        # individually with PSM 7 (single line) — far more accurate than
        # whole-image OCR and completely bypasses the coloured-header problem.
        import cv2 as _cv2
        import numpy as _np

        _gray_for_grid = _np.array(img.convert("L"))
        # Case A candidate: full drawn grid → per-cell OCR.  Collected (not
        # returned early) so it competes with the other strategies on score.
        _grid_rows = None
        _grid = _detect_table_grid(_gray_for_grid)
        if _grid is not None:
            _col_bounds, _row_bounds = _grid
            _grid_rows = _extract_cells_from_grid(img, _col_bounds, _row_bounds) or None

        # ── Coloured header band (e.g. purple transaction header) ─────────────
        # The band's word positions define exact column boundaries — far more
        # reliable than inferring columns from transaction-row word gaps.
        _band = _detect_header_band(_gray_for_grid)
        _header_col_bounds = None
        _header_labels = None
        _table_y_start = None
        if _band is not None:
            _bt, _bb = _band
            _hb = _columns_from_header_band(img, _bt, _bb)
            if _hb is not None:
                _header_col_bounds, _header_labels = _hb
                _table_y_start = _bb   # transaction rows start just below the band

        # ── Preprocessing (lesson from professional converters) ───────────────
        # 1. Scale up small images — Tesseract accuracy drops below ~150 DPI.
        # 2. Adaptive binarisation: handles dark/coloured header rows (purple,
        #    green, etc.) and uneven lighting without a second inverted-image pass.
        # 3. Use LSTM engine (--oem 3) + uniform-block PSM (--psm 6).
        gray = _np.array(img.convert("L"))

        # Scale up if the image is small
        if img.width < 1500:
            scale = max(2, 1500 // img.width)
            gray  = _cv2.resize(gray, (img.width * scale, img.height * scale),
                                interpolation=_cv2.INTER_LANCZOS4)

        # Auto-handle light-text-on-dark-background images (e.g. dark-mode
        # screenshots).  Tesseract reads black-on-white, so we invert the
        # grayscale when the text is lighter than the background.
        # Polarity via Otsu: the minority pixel class is the "ink" (text).  If
        # dark pixels are the MAJORITY, the background is dark → invert.  This
        # correctly leaves dark-text-on-medium-grey docs (e.g. tinted bank
        # statements) untouched, unlike a naive mean-brightness threshold.
        _otsu_thr, _ = _cv2.threshold(gray, 0, 255, _cv2.THRESH_BINARY + _cv2.THRESH_OTSU)
        if int((gray < _otsu_thr).sum()) >= int((gray >= _otsu_thr).sum()):
            gray = 255 - gray

        from PIL import Image as _PILImage

        def _mean_conf(r):
            cs = [int(c) for i, c in enumerate(r["conf"])
                  if r["text"][i].strip() and int(c) >= 0]
            return (sum(cs) / len(cs)) if cs else 0.0

        # Adaptive threshold: each 31×31 local region binarised independently
        bw = _cv2.adaptiveThreshold(
            gray, 255,
            _cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            _cv2.THRESH_BINARY, 31, 10,
        )
        preprocessed = _PILImage.fromarray(bw)
        raw = pytesseract.image_to_data(
            preprocessed, output_type=Output.DICT, config="--oem 3 --psm 6"
        )

        # Low confidence ⇒ the local adaptive threshold is producing speckle
        # (typical of low-contrast camera photos of screens).  Retry with a
        # denoise + global-Otsu pass and keep whichever reads more confidently.
        # Clean screenshots score 70-95 here and never trigger this, so their
        # behaviour is unchanged.
        if _mean_conf(raw) < 55:
            den = _cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            _, bw2 = _cv2.threshold(den, 0, 255, _cv2.THRESH_BINARY + _cv2.THRESH_OTSU)
            pre2 = _PILImage.fromarray(bw2)
            raw2 = pytesseract.image_to_data(
                pre2, output_type=Output.DICT, config="--oem 3 --psm 6"
            )
            if _mean_conf(raw2) > _mean_conf(raw):
                preprocessed, raw = pre2, raw2

        _tess_conf = _mean_conf(raw)   # used below to decide on the RapidOCR fallback

        scale_x = img.width  / preprocessed.width
        scale_y = img.height / preprocessed.height

        # Remap coordinates back to original image space
        data = {
            "text": raw["text"],
            "conf": raw["conf"],
            "left":   [int(v * scale_x) for v in raw["left"]],
            "top":    [int(v * scale_y) for v in raw["top"]],
            "width":  [int(v * scale_x) for v in raw["width"]],
            "height": [int(v * scale_y) for v in raw["height"]],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR failed: {str(e)}")

    # ── Build word list with y-band line grouping ────────────────────────────
    # For photos, pytesseract's block/line numbering can misalign words on the
    # same visual row.  Instead, cluster words by y-centre within ±50% of the
    # average character height.
    from collections import defaultdict

    all_words: list[dict] = []
    for i in range(len(data["text"])):
        text = data["text"][i].strip()
        conf = int(data["conf"][i])
        # Keep every word Tesseract actually emitted (conf >= 0); drop only the
        # -1 structural markers.  Tesseract confidence is unreliable — it assigns
        # conf 0-15 to perfectly-correct words (e.g. "CHECK 1249", "SERVICE
        # CHARGE"), so a conf>20 gate silently DROPS valid table cells.  Garbage
        # is handled downstream by the best-of-N scorer, not by this gate.
        if not text or conf < 0:
            continue
        text = _desplit_glued(text)   # 'POSPURCHASE' → 'POS PURCHASE'
        w = data["width"][i]
        h = data["height"][i]
        all_words.append({
            "text":   text,
            "left":   data["left"][i],
            "top":    data["top"][i],
            "right":  data["left"][i] + w,
            "width":  w,
            "height": max(h, 1),
        })

    # Low Tesseract confidence ⇒ a degraded photo (e.g. a camera shot of a screen).
    # Swap in RapidOCR — a free, local neural OCR that reads such images far better —
    # if it's installed. Clean docs score 70-95 here and keep their tuned Tesseract
    # output untouched; only the hard photos get the heavier neural engine.
    if _tess_conf < 65:
        rapid_words = _ocr_words_rapidocr(img)
        if rapid_words:
            log.info("Low Tesseract confidence (%.0f) — using RapidOCR (%d regions).",
                     _tess_conf, len(rapid_words))
            all_words = rapid_words

    wb = openpyxl.Workbook()
    ws = wb.active

    if not all_words:
        write_data_to_sheet(ws, [["Result"], ["No text could be extracted via OCR."]], "OCR Result")
        return save_workbook(wb, stem)

    # Y-band clustering: group words whose vertical centres are within
    # tolerance of the current running line centre.
    avg_h   = sum(w["height"] for w in all_words) / len(all_words)
    y_tol   = avg_h * 0.6
    sorted_w = sorted(all_words, key=lambda w: (w["top"] + w["height"] / 2, w["left"]))

    line_groups: list[list[dict]] = []
    cur_line:    list[dict]       = [sorted_w[0]]
    cur_cy = sorted_w[0]["top"] + sorted_w[0]["height"] / 2

    for wd in sorted_w[1:]:
        wd_cy = wd["top"] + wd["height"] / 2
        if abs(wd_cy - cur_cy) <= y_tol:
            cur_line.append(wd)
            cur_cy = sum(w["top"] + w["height"] / 2 for w in cur_line) / len(cur_line)
        else:
            line_groups.append(sorted(cur_line, key=lambda w: w["left"]))
            cur_line = [wd]
            cur_cy   = wd_cy
    line_groups.append(sorted(cur_line, key=lambda w: w["left"]))

    img_width = img.width

    # ── Word-based column detection cascade (unchanged logic) ─────────────────
    # Computes column boundaries from a set of line_groups, ignoring any header
    # band.  This is the "Case C" detector and is reused as the safe baseline.
    def _cascade_col_boundaries(lgs):
        col_boundaries = None

        # 1. Header-keyword based (bank statements etc.)
        col_boundaries = _find_col_boundaries(lgs, img_width)

        # 2. Zone histogram: divide width into 50 coarse zones.
        if col_boundaries is None:
            n_zones  = 50
            zone_w   = max(1.0, img_width / n_zones)
            zone_cov = [False] * n_zones
            for wd in all_words:
                z0 = int(wd["left"]  / zone_w)
                z1 = int(wd["right"] / zone_w)
                for z in range(max(0, z0), min(n_zones, z1 + 1)):
                    zone_cov[z] = True
            min_gap_zones = 2
            histo_bounds  = [0]
            z = 0
            while z < n_zones:
                if not zone_cov[z]:
                    gap_start = z
                    while z < n_zones and not zone_cov[z]:
                        z += 1
                    gap_end = z
                    if gap_end - gap_start >= min_gap_zones:
                        mid_px = int(((gap_start + gap_end) / 2) * zone_w)
                        histo_bounds.append(mid_px)
                else:
                    z += 1
            histo_bounds.append(img_width)
            if len(histo_bounds) > 2:
                col_boundaries = histo_bounds

        # 3. Right-edge clustering (financial tables right-align amounts).
        if col_boundaries is None:
            _NUM_CLEAN_RE = _re.compile(r'[\$\d,\.]+')
            _NUM_RE = _re.compile(r'^\$?[\d,]+\.\d{2}$')

            def _clean_num(text: str) -> str:
                m = _NUM_CLEAN_RE.search(text)
                return m.group(0).rstrip('.') if m else ""

            right_edges: list[int] = []
            for line in lgs:
                for wd in line:
                    if _NUM_RE.match(_clean_num(wd["text"])):
                        right_edges.append(wd["right"])
            if len(right_edges) >= 4:
                right_edges.sort()
                re_clusters: list[list[int]] = [[right_edges[0]]]
                for r in right_edges[1:]:
                    if r - re_clusters[-1][-1] <= 15:
                        re_clusters[-1].append(r)
                    else:
                        re_clusters.append([r])
                sig = [c for c in re_clusters if len(c) >= 2]
                if len(sig) >= 2:
                    centres = sorted(sum(c) // len(c) for c in sig)
                    first_num_left = min(
                        (wd["left"] for line in lgs for wd in line
                         if _NUM_RE.match(_clean_num(wd["text"]))),
                        default=centres[0] - 50,
                    )
                    bounds = [0, max(5, first_num_left - 5)]
                    for i in range(len(centres) - 1):
                        bounds.append((centres[i] + centres[i + 1]) // 2)
                    bounds.append(img_width)
                    if len(bounds) >= 4:
                        col_boundaries = bounds

        # 4. Gap-distribution fallback
        if col_boundaries is None:
            col_boundaries = _detect_columns_from_gaps(lgs, img_width)

        return col_boundaries

    # ── Render rows from given boundaries (unchanged logic) ───────────────────
    # table_y_start: if set, rows above it are emitted as plain metadata text
    # rows above the table (used by the coloured-header-band strategy).
    def _render_rows(col_boundaries, header_labels, table_y_start):
        meta_lgs = []
        table_lgs = line_groups
        if table_y_start is not None:
            meta_lgs  = [lg for lg in line_groups if lg[0]["top"] <  table_y_start]
            table_lgs = [lg for lg in line_groups if lg[0]["top"] >= table_y_start]

        if col_boundaries and len(col_boundaries) > 2:
            rows = []
            if header_labels and len(header_labels) == len(col_boundaries) - 1:
                rows.append(list(header_labels))
            for line in table_lgs:
                cells = _line_to_cells_fixed(line, col_boundaries)
                if any(c.strip() for c in cells):
                    rows.append(cells)
            if not rows:
                rows = [["Result"], ["No structured content detected."]]
        else:
            rows = [["Line #", "Extracted Text"]]
            for i, line in enumerate(table_lgs, start=1):
                text = " ".join(wd["text"] for wd in line)
                rows.append([i, text])

        if meta_lgs:
            meta_rows = []
            for line in meta_lgs:
                text = " ".join(wd["text"] for wd in line).strip()
                if text:
                    meta_rows.append([text])
            rows = meta_rows + [[""]] + rows
        return rows

    # ── Collect candidate results from each applicable strategy ───────────────
    candidates: list[tuple[str, list]] = []

    # Case A: drawn grid
    if _grid_rows:
        candidates.append(("grid", _grid_rows))

    # Case B: coloured header band gives exact column boundaries
    if _header_col_bounds is not None:
        candidates.append((
            "header_band",
            _render_rows(_header_col_bounds, _header_labels, _table_y_start),
        ))

    # Case C: word-based cascade (always present — the safe baseline)
    candidates.append(("word_based", _render_rows(_cascade_col_boundaries(line_groups), None, None)))

    # Case D: left-edge clustering (best for left-aligned multi-column tables)
    _left_bounds = _columns_from_left_edges(line_groups, img_width)
    if _left_bounds is not None:
        candidates.append(("left_aligned", _render_rows(_left_bounds, None, None)))

    # Case E: faint vertical rules → columns (bordered light-bg statements, e.g.
    # SBI). Uses the ruled column x-boundaries with the good whole-image OCR.
    # Also recover the table's header row (which full-page OCR usually skips) by
    # OCRing the header band on its own, and place metadata above it.
    _ruled_bounds = _columns_from_ruled_lines(_gray_for_grid, img_width)
    if _ruled_bounds is not None:
        _ruled_hdr = _header_row_for_ruled(img, _gray_for_grid, _ruled_bounds)
        _hdr_cells, _hdr_top = _ruled_hdr if _ruled_hdr else (None, None)
        # Render every line into the ruled columns; insert the recovered header
        # row at the table boundary (full-page OCR skips it, so no line_group
        # exists there).  Metadata stays rendered in columns (scores best).
        _rrows: list[list[str]] = []
        _hdr_done = False
        for ln in line_groups:
            if _hdr_cells and not _hdr_done and ln[0]["top"] >= _hdr_top:
                _rrows.append(list(_hdr_cells))
                _hdr_done = True
            cells = _line_to_cells_fixed(ln, _ruled_bounds)
            if any(c.strip() for c in cells):
                _rrows.append(cells)
        if _hdr_cells and not _hdr_done:
            _rrows.append(list(_hdr_cells))
        candidates.append(("ruled_columns", _rrows))

    # ── Pick the best-scoring candidate ───────────────────────────────────────
    # Confidence priors: a detected coloured-header band is a strong image-type
    # signal, so it gets a bonus (it also legitimately splits metadata text rows
    # above the table, which would otherwise look "inconsistent" to the scorer).
    _CONF_BONUS = {"header_band": 2.0}

    def _cand_score(c):
        return _score_table_rows(c[1]) + _CONF_BONUS.get(c[0], 0.0)

    best_name, rows = max(candidates, key=_cand_score)
    log.info(
        "Image strategy scores: %s → chose '%s'",
        {n: round(_cand_score((n, r)), 2) for n, r in candidates},
        best_name,
    )

    # Tilt refinement: if the chosen rows glued two phone numbers into one cell,
    # the photo is tilted enough to have smeared the row grouping — re-group at a
    # few tilt angles and keep the cleanest. No-op unless that defect is present.
    rows = _refine_tilted_rows(rows, all_words, img_width, y_tol)

    # Final validation pass: drop Excel UI chrome when the image is a screenshot
    # of a spreadsheet (no-op for ordinary photos/scans).
    rows = _strip_spreadsheet_chrome(rows)

    # Split a column where OCR glued names onto their phone numbers (no-op unless a
    # column is predominantly "name 9876543210").
    rows = _split_glued_name_phone_columns(rows)

    if merge_cols:
        rows = _merge_repeating_cols(rows)

    write_data_to_sheet(ws, rows, "OCR Result")
    return save_workbook(wb, stem)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/api")
async def api_root():
    """Health/info endpoint (also used as the Render health check)."""
    return {"message": "DocToExcel API is running. POST /api/convert to convert files."}


@app.get("/api/cloud-ocr-status")
async def cloud_ocr_status():
    """Which cloud OCR engine is configured server-side ('gemini' | 'azure' | null).
    Used by the UI to show the admin-mode 'AI OCR (cloud)' toggle's real availability."""
    return {"engine": _cloud_ocr_engine()}


def _run_conversion(params: dict) -> tuple[str, str]:
    """
    Synchronous conversion core — runs inside a worker thread.
    Returns (output_filename, log_message). Raises HTTPException(423) when a PDF
    needs a password; other failures bubble up as exceptions.
    """
    content       = params["content"]
    ext           = params["ext"]
    stem          = params["stem"]
    original_name = params["original_name"]
    mode          = params["mode"]
    password      = params["password"]
    merge_cols    = params["merge_cols"]
    use_azure     = params["use_azure"]

    if ext == ".csv":
        output_filename = convert_csv(content, stem)
        msg = f'"{original_name}" has been converted successfully! Your Excel file is ready.'
    elif ext == ".txt":
        output_filename = convert_txt(content, stem)
        msg = f'"{original_name}" has been converted successfully! Your Excel file is ready.'
    elif ext == ".pdf":
        output_filename = convert_pdf(content, stem, mode=mode, password=password)
        if mode == "single":
            msg = f'"{original_name}" converted! All tables are combined in one sheet.'
        else:
            msg = f'"{original_name}" converted! Each table/section is on a separate sheet.'
    elif ext in (".docx", ".doc"):
        output_filename = convert_docx(content, stem)
        msg = f'"{original_name}" converted! Tables and text have been extracted to Excel.'
    elif ext in (".png", ".jpg", ".jpeg"):
        do_merge = merge_cols.lower() == "true"
        want_cloud = use_azure.lower() == "true"   # admin "AI OCR (cloud)" toggle
        output_filename = None
        cloud_err = None                            # reason the cloud attempt failed
        engine = _cloud_ocr_engine() if want_cloud else None
        if engine == "gemini":
            try:
                output_filename = convert_image_gemini(content, stem, mode=mode, merge_cols=do_merge)
                msg = f'"{original_name}" converted via Google Gemini.'
            except CloudOCRUnavailable as e:
                cloud_err = str(e)
                log.warning("Gemini path unavailable, falling back to Tesseract: %s", e)
        elif engine == "azure":
            try:
                output_filename = convert_image_azure(content, stem, mode=mode, merge_cols=do_merge)
                msg = f'"{original_name}" converted via Azure Document Intelligence.'
            except CloudOCRUnavailable as e:
                cloud_err = str(e)
                log.warning("Azure path unavailable, falling back to Tesseract: %s", e)
        if output_filename is None:
            output_filename = convert_image(content, stem, original_name, merge_cols=do_merge)
            if not want_cloud:
                msg = f'"{original_name}" processed via OCR. Extracted text has been placed in the Excel file.'
            elif engine is None:
                msg = (f'"{original_name}" processed via Tesseract OCR '
                       f'(no cloud engine configured — see README).')
            elif cloud_err and any(t in cloud_err.lower()
                                   for t in ("429", "quota", "limit: 0", "resource_exhausted", "rate limit")):
                msg = (f'"{original_name}" processed via Tesseract OCR '
                       f"(cloud OCR daily quota reached — try again later).")
            else:
                msg = (f'"{original_name}" processed via Tesseract OCR '
                       f'(cloud OCR temporarily unavailable).')
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    log.info("Conversion result: %s", msg)
    return output_filename, msg


@app.post("/api/convert")
async def convert_file(file: UploadFile = File(...), mode: str = Form("single"), password: str = Form(""), merge_cols: str = Form("false"), use_azure: str = Form("false")):
    """
    Accept an uploaded file and enqueue it for conversion.
    Returns immediately with { job_id, status, position } — poll
    GET /api/convert/{job_id} for the result, or DELETE it to cancel while queued.
    """
    original_name = file.filename or "document"
    path = Path(original_name)
    ext = path.suffix.lower()
    stem = path.stem[:40]  # cap length for filename

    log.info("Received file: %s (ext=%s, content_type=%s)", original_name, ext, file.content_type)

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f'File type "{ext}" is not supported. Please upload PDF, DOCX, PNG, JPG, CSV, or TXT.',
        )

    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")

    _prune_jobs()
    job = Job({
        "content": content,
        "ext": ext,
        "stem": stem,
        "original_name": original_name,
        "mode": mode,
        "password": password,
        "merge_cols": merge_cols,
        "use_azure": use_azure,
    })
    JOBS[job.id] = job
    job.task = asyncio.create_task(_process_job(job))

    return JSONResponse(_job_view(job))


@app.get("/api/convert/{job_id}")
async def convert_status(job_id: str):
    """Poll the status (and result, once done) of a queued/running conversion."""
    job = JOBS.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found or expired.")
    return JSONResponse(_job_view(job))


@app.delete("/api/convert/{job_id}")
async def cancel_conversion(job_id: str):
    """
    Cancel a conversion — only allowed while it is still queued. A job that has
    already started processing can't be safely interrupted.
    """
    job = JOBS.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found or expired.")

    if job.status == "queued":
        job.status = "cancelled"        # set first so the worker skips it if it wins the slot race
        if job.task:
            job.task.cancel()
        log.info("Cancelled queued job %s", job_id)
        return JSONResponse({"job_id": job.id, "status": "cancelled"})

    return JSONResponse(
        {"job_id": job.id, "status": job.status,
         "message": "Cannot cancel — conversion already processing or finished."},
        status_code=409,
    )


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """Serve a previously converted Excel file."""
    # Security: ensure no path traversal
    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid filename.")

    file_path = OUTPUTS_DIR / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f'File "{safe_name}" not found.')

    return FileResponse(
        path=file_path,
        filename=safe_name,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/api/files")
async def list_files():
    """List all converted files in the outputs directory."""
    files = []
    for f in sorted(OUTPUTS_DIR.glob("*.xlsx"), key=lambda x: x.stat().st_mtime, reverse=True):
        files.append({
            "filename": f.name,
            "size_kb": round(f.stat().st_size / 1024, 1),
        })
    return {"files": files}


# ── Serve the built frontend (production) ─────────────────────────────────────
#
# In production (Docker / Render) the React app is built to frontend/dist and
# served by this same server, so the whole app is one origin (no CORS, one
# service). In local dev this directory doesn't exist — run the Vite dev server
# separately (it proxies /api/* back here), and this mount is simply skipped.
#
# Mounted LAST so every /api/* route above takes precedence over static files.

_FRONTEND_DIST = Path(__file__).resolve().parent.parent / "frontend" / "dist"
if _FRONTEND_DIST.is_dir():
    app.mount("/", StaticFiles(directory=str(_FRONTEND_DIST), html=True), name="frontend")
    log.info("Serving frontend from %s", _FRONTEND_DIST)
else:
    @app.get("/")
    async def root():
        return {"message": "DocToExcel API is running. Frontend not built; run the Vite dev server."}
